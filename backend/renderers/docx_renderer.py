"""Markdown -> python-docx rendering helpers.

The policy generation / migration pipeline sometimes hands the docx builder
markdown-flavored prose (`## Heading`, `**bold**`, `* bullet`, `[link](url)`, …).
python-docx writes whatever string it's given into a single Normal-style run,
so these symbols leak into the .docx as literal characters.

This module renders that markdown into real Word formatting:
- ATX headings (`#` .. `######`) -> Heading 1..6 styles
- Unordered list lines (`*` / `-` / `+` prefix) -> "List Bullet" paragraphs
- Ordered list lines (`1. ` etc) -> "List Number" paragraphs
- Blockquote lines (`> `) -> "Intense Quote" paragraphs
- Inline `**bold**` / `__bold__` -> bold runs
- Inline `*italic*` / `_italic_` -> italic runs
- Inline `` `code` `` -> Consolas monospace runs
- Inline `~~strike~~` -> strikethrough runs
- Inline `[text](url)` -> hyperlink runs (text only; the URL is dropped onto
  the run with a styled appearance — python-docx hyperlink XML is overkill
  for cosmetic rendering and we don't have a target field for the URL).

This is a "good enough" CommonMark subset — not a full parser. The goal is
to render the markdown the LLM/migration source emits in practice, not to
implement a complete spec.
"""
from __future__ import annotations

import re
from typing import Iterable, TYPE_CHECKING

if TYPE_CHECKING:  # pragma: no cover
    from docx.document import Document
    from docx.text.paragraph import Paragraph


# ── Block-level patterns ─────────────────────────────────────────────────────

_ATX_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_UNORDERED_LIST_RE = re.compile(r"^([*\-+])\s+(.+)$")
_ORDERED_LIST_RE = re.compile(r"^(\d+)[.)]\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")

# Inline patterns. Ordering matters - run them in this sequence on each
# paragraph so the longer markers (`**`, `***`, `~~`) win over `*`/`_`.
_INLINE_RULES: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"\*\*\*(.+?)\*\*\*"), "bolditalic"),
    (re.compile(r"___(.+?)___"), "bolditalic"),
    (re.compile(r"\*\*(.+?)\*\*"), "bold"),
    (re.compile(r"__(.+?)__"), "bold"),
    (re.compile(r"~~(.+?)~~"), "strike"),
    (re.compile(r"`([^`]+)`"), "code"),
    # Italic with single * or _ — only match when the marker is NOT part of
    # a longer ** / __ sequence. The longer-marker rules above are run first
    # and consume those, so a bare * here is genuinely italic.
    (re.compile(r"(?<!\*)\*(?!\*)([^*\n]+?)(?<!\*)\*(?!\*)"), "italic"),
    (re.compile(r"(?<!_)_(?!_)([^_\n]+?)(?<!_)_(?!_)"), "italic"),
    # Markdown link [text](url). We keep the visible text and drop the URL
    # into a trailing parenthetical so the destination is still visible.
    (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), "link"),
]


def _split_inline(text: str) -> list[tuple[str, str]]:
    """Tokenize a single paragraph of markdown into (style, fragment) runs.

    Styles are: "plain", "bold", "italic", "bolditalic", "code", "strike",
    "link". The list of fragments concatenated equals the original visible
    text (with markdown markers stripped).
    """
    if not text:
        return [("plain", "")]

    # Use placeholder substitution: replace matched spans with sentinels we
    # can split on later. This keeps the linear-walk simple and robust to
    # adjacent or nested formatting.
    sentinels: list[tuple[str, str]] = []  # (style, fragment)
    work = text

    def _stash(style: str) -> str:
        idx = len(sentinels)
        sentinels.append((style, ""))  # placeholder, filled in below
        return f"\x00SENT{idx}\x00"

    for pattern, style in _INLINE_RULES:
        def _sub(match: re.Match[str]) -> str:
            token = _stash(style)
            # For markdown links, surface the URL alongside the visible text.
            if style == "link":
                visible = match.group(1)
                url = match.group(2)
                fragment = f"{visible} ({url})"
            else:
                fragment = match.group(1)
            sentinels[-1] = (style, fragment)
            return token

        work = pattern.sub(_sub, work)

    # Reassemble: split on sentinels, interleave plain segments with stashed
    # formatted runs.
    tokens = re.split(r"\x00SENT(\d+)\x00", work)
    runs: list[tuple[str, str]] = []
    for i, chunk in enumerate(tokens):
        if i % 2 == 0:
            if chunk:
                runs.append(("plain", chunk))
        else:
            style, fragment = sentinels[int(chunk)]
            runs.append((style, fragment))
    return runs or [("plain", "")]


def _apply_run_style(run, style: str) -> None:
    if style == "bold":
        run.bold = True
    elif style == "italic":
        run.italic = True
    elif style == "bolditalic":
        run.bold = True
        run.italic = True
    elif style == "strike":
        run.font.strike = True
    elif style == "code":
        run.font.name = "Consolas"
    elif style == "link":
        # Visual treatment only; python-docx hyperlink XML is heavier than
        # what we need for cosmetic rendering. The URL is appended in the
        # fragment text by _split_inline.
        run.font.color.rgb = _LINK_COLOR
        run.underline = True


def _emit_paragraph_with_inline(paragraph, text: str) -> None:
    for style, fragment in _split_inline(text):
        if not fragment:
            continue
        run = paragraph.add_run(fragment)
        if style != "plain":
            _apply_run_style(run, style)


# ── Public entry points ──────────────────────────────────────────────────────

# Bright FastAPI-ish blue, just for visual cue on links. RGB tuple kept here
# so callers don't have to import docx internals.
from docx.shared import RGBColor  # noqa: E402  (intentional after constants)

_LINK_COLOR = RGBColor(0x1A, 0x73, 0xE8)


def render_markdown_into(doc, text: str, *, default_style: str | None = None) -> None:
    """Append the markdown `text` into `doc` as properly-styled paragraphs.

    Args:
        doc: a python-docx Document.
        text: markdown body. May contain multiple paragraphs separated by
            blank lines. Treated as already-decoded str.
        default_style: optional paragraph style for plain prose. When None,
            python-docx's "Normal" is used.
    """
    if text is None:
        return
    text = str(text)
    if not text.strip():
        return

    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        stripped = line.lstrip()

        # ATX heading
        h = _ATX_HEADING_RE.match(stripped)
        if h:
            level = min(len(h.group(1)), 6)
            heading_text = h.group(2).strip()
            paragraph = doc.add_heading("", level=level)
            _emit_paragraph_with_inline(paragraph, heading_text)
            continue

        # Unordered list
        u = _UNORDERED_LIST_RE.match(stripped)
        if u:
            paragraph = doc.add_paragraph(style="List Bullet")
            _emit_paragraph_with_inline(paragraph, u.group(2).strip())
            continue

        # Ordered list
        o = _ORDERED_LIST_RE.match(stripped)
        if o:
            paragraph = doc.add_paragraph(style="List Number")
            _emit_paragraph_with_inline(paragraph, o.group(2).strip())
            continue

        # Blockquote
        b = _BLOCKQUOTE_RE.match(stripped)
        if b:
            quoted = b.group(1).strip()
            try:
                paragraph = doc.add_paragraph(style="Intense Quote")
            except (KeyError, ValueError):
                # Style not available in the current docx template; fall back.
                paragraph = doc.add_paragraph()
            _emit_paragraph_with_inline(paragraph, quoted)
            continue

        # Blank line -> empty paragraph (mirrors source spacing without
        # collapsing it onto the previous paragraph's run).
        if not stripped:
            doc.add_paragraph()
            continue

        # Plain paragraph (with inline formatting)
        if default_style:
            try:
                paragraph = doc.add_paragraph(style=default_style)
            except (KeyError, ValueError):
                paragraph = doc.add_paragraph()
        else:
            paragraph = doc.add_paragraph()
        _emit_paragraph_with_inline(paragraph, stripped)


def render_markdown_bullet(doc, text: str) -> None:
    """Append `text` as a single bullet paragraph with inline markdown
    formatting applied. Used for list-of-strings content where the source
    has chosen a bullet semantically and we only need to style the inline
    runs."""
    if text is None:
        return
    paragraph = doc.add_paragraph(style="List Bullet")
    _emit_paragraph_with_inline(paragraph, str(text).strip())
