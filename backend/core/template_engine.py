"""
Midnight Core — template_engine.py
Takeoff LLC

Real template engine: doc_type + variant → a styled DOCX shell from
backend/templates/packs/ (36 shells: 9 categories × 4 variants, imported from
the midnight-template-build collection). The shell's body is cleared but its
styles, fonts, headers, and section setup survive — so everything rendered
into it (headings, body, bullets) inherits the pack's typography.

Fail-open by design: a missing pack or unreadable shell falls back to a blank
python-docx Document, which is exactly the pre-template-engine behavior.
"""
from __future__ import annotations

import copy
import logging
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger("midnight.template_engine")

PACKS_DIR = Path(__file__).resolve().parent.parent / "templates" / "packs"

VARIANTS = ("formal", "modern", "detailed", "executive")
DEFAULT_VARIANT = "modern"

# Studio doc types → pack category directories.
DOC_TYPE_TO_CATEGORY = {
    "POLICY": "policy",
    "STANDARD": "standard",
    "SOP": "procedure",
    "PROCEDURE": "procedure",
    "INCIDENT_RUNBOOK": "incident_runbook",
    "PLAYBOOK": "incident_runbook",
    "PROCESS_FLOW": "process_flow",
    "TRAINING": "training",
    "TRAINING_MODULE": "training",
    "RISK_ASSESSMENT": "risk_assessment",
    "AUDIT_PACKAGE": "audit_package",
    "AI_GOVERNANCE": "ai_governance",
}

# Placeholder tokens the shells carry in headers/footers.
_PLACEHOLDERS = ("{{CLASSIFICATION}}", "{{ORG}}", "{{ORGANIZATION}}", "{{TITLE}}")


def normalize_variant(variant: str | None) -> str:
    v = str(variant or "").strip().lower()
    return v if v in VARIANTS else DEFAULT_VARIANT


def shell_path(doc_type: str, variant: str | None = None) -> Path | None:
    category = DOC_TYPE_TO_CATEGORY.get(str(doc_type or "").strip().upper())
    if not category:
        return None
    v = normalize_variant(variant)
    path = PACKS_DIR / category / v / f"{category}_{v}.docx"
    return path if path.exists() else None


_BUILTIN_STYLE_NAMES = (
    "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "List Paragraph", "Strong", "Quote",
)


def _normalize_builtin_style_names(doc) -> None:
    """Pandoc-built shells store UI style names ("Heading 1") where python-docx
    expects internal names ("heading 1"), so styles["Heading 1"] — and
    doc.add_heading — KeyError even though iteration shows the style.
    Reassigning the name through python-docx's setter writes the internal
    form and makes lookups (and rendering) work."""
    from docx.styles import BabelFish

    for style in doc.styles:
        if style.name in _BUILTIN_STYLE_NAMES:
            try:
                # Write the internal form (e.g. "heading 1") — the getter
                # translates it back to the UI form, and lookups resolve.
                style.name = BabelFish.ui2internal(style.name)
            except Exception:
                continue


def _ensure_table_grid_style(doc) -> None:
    """The render path styles metadata tables as "Table Grid"; pandoc shells
    don't define it. Register a minimal bordered table style under that name
    so table.style = "Table Grid" resolves."""
    from docx.enum.style import WD_STYLE_TYPE

    try:
        doc.styles["Table Grid"]
        return
    except KeyError:
        pass
    try:
        style = doc.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE, builtin=True)
        # Single-line borders on all edges, matching Word's built-in look.
        from docx.oxml.parser import parse_xml
        from docx.oxml.ns import nsdecls

        edges = "".join(
            f'<w:{edge} w:val="single" w:sz="4" w:color="auto"/>'
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV")
        )
        tbl_pr = parse_xml(f"<w:tblPr {nsdecls('w')}><w:tblBorders>{edges}</w:tblBorders></w:tblPr>")
        style.element.append(tbl_pr)
    except Exception as exc:
        logger.warning("could not register Table Grid style: %s", exc)


def _ensure_paragraph_styles(doc) -> None:
    """Register paragraph styles the render path uses that pandoc shells lack
    (List Bullet crashed a live export; List Number is the same class)."""
    from docx.enum.style import WD_STYLE_TYPE

    for name in ("List Bullet", "List Number"):
        try:
            doc.styles[name]
            continue
        except KeyError:
            pass
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=True)
            try:
                style.base_style = doc.styles["List Paragraph"]
            except KeyError:
                pass
        except Exception as exc:
            logger.warning("could not register %s style: %s", name, exc)


def _clear_body(doc) -> None:
    """Remove all body content, keeping section properties (headers/margins)."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _replace_placeholder_text(doc, replacements: dict[str, str]) -> None:
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for token, value in replacements.items():
                    if token in paragraph.text:
                        for run in paragraph.runs:
                            if token in run.text:
                                run.text = run.text.replace(token, value)
                        # Token may span runs; last resort rebuilds the text.
                        if token in paragraph.text:
                            paragraph.text = paragraph.text.replace(token, value)


def _parse_hex_color(value: str | None) -> RGBColor | None:
    raw = str(value or "").strip().lstrip("#")
    if len(raw) != 6:
        return None
    try:
        return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except ValueError:
        return None


def apply_branding(doc, *, organization: str = "", title: str = "", primary_color: str | None = None,
                   classification: str = "Internal") -> None:
    """Tenant branding: header placeholders + brand color on Title/Heading styles."""
    _replace_placeholder_text(doc, {
        "{{CLASSIFICATION}}": classification or "Internal",
        "{{ORG}}": organization or "",
        "{{ORGANIZATION}}": organization or "",
        "{{TITLE}}": title or "",
    })
    color = _parse_hex_color(primary_color)
    if color is not None:
        for style_name in ("Title", "Heading 1", "Heading 2"):
            try:
                doc.styles[style_name].font.color.rgb = color
            except (KeyError, AttributeError):
                continue


def apply_logo(doc, logo_url: str | None) -> None:
    """Insert the tenant logo into the document header. Fail-open."""
    url = str(logo_url or "").strip()
    if not url.lower().startswith(("http://", "https://")):
        return
    logo = _fetch_logo_bytes(url)
    if logo:
        _insert_logo(doc, logo)


def preview_path(category: str, variant: str) -> Path | None:
    """Validated path to a pack's preview PNG — inputs checked against the
    known category/variant sets, so no traversal is possible."""
    cat = str(category or "").strip().lower()
    v = str(variant or "").strip().lower()
    if cat not in set(DOC_TYPE_TO_CATEGORY.values()) or v not in VARIANTS:
        return None
    path = PACKS_DIR / cat / v / f"{cat}_{v}_preview.png"
    return path if path.exists() else None


def _fetch_logo_bytes(url: str) -> bytes | None:
    """Fetch a tenant logo, fail-open. Small timeout — rendering must not hang."""
    try:
        import httpx

        resp = httpx.get(url, timeout=5.0, follow_redirects=True)
        resp.raise_for_status()
        if len(resp.content) > 2 * 1024 * 1024:
            return None
        return resp.content
    except Exception as exc:
        logger.warning("logo fetch failed (%s); rendering without logo", exc)
        return None


def _insert_logo(doc, logo_bytes: bytes) -> None:
    from io import BytesIO

    from docx.shared import Inches

    try:
        header = doc.sections[0].header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(BytesIO(logo_bytes), width=Inches(1.1))
    except Exception as exc:
        logger.warning("logo insert failed (%s); rendering without logo", exc)


def load_template_shell(doc_type: str, variant: str | None = None, *, branding: dict | None = None):
    """Return a python-docx Document ready to render into.

    Styled shell when a pack exists for (doc_type, variant); blank Document
    otherwise. Never raises — the export path must not fail on styling.
    """
    path = shell_path(doc_type, variant)
    if path is None:
        doc = Document()
        doc.is_template_shell = False
        return doc
    try:
        doc = Document(str(path))
        _normalize_builtin_style_names(doc)
        _ensure_table_grid_style(doc)
        _ensure_paragraph_styles(doc)
        _clear_body(doc)
        doc.is_template_shell = True
        if branding:
            apply_branding(
                doc,
                organization=str(branding.get("organization") or ""),
                title=str(branding.get("title") or ""),
                primary_color=branding.get("primary_color"),
                classification=str(branding.get("classification") or "Internal"),
            )
            apply_logo(doc, branding.get("logo_url"))
        return doc
    except Exception as exc:
        logger.warning("template shell %s unusable (%s); falling back to blank", path, exc)
        doc = Document()
        doc.is_template_shell = False
        return doc
