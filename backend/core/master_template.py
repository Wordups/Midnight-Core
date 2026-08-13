"""
Midnight Core — master_template.py

A single, well-designed, reusable Word template built programmatically (so it's
parametric and never drifts from a binary file). It replaces the flat pandoc
shells with real typographic hierarchy and, crucially, produces a document that
is clean to edit locally in Word:

  * every visual decision lives in a NAMED STYLE (Title, Heading 1/2/3, Normal,
    plus a few custom cover styles), so content the user adds later inherits the
    design instead of looking pasted-in;
  * the Table of Contents is a live Word field — right-click > Update Field
    repaginates it after edits;
  * page numbers are a footer field;
  * fonts are Office-standard (Segoe UI / Calibri), so nothing renders as a
    missing-font fallback on the user's machine.

It emits the same {{TOKEN}} placeholders the fill path already knows, and NO
sample body sections — `_build_docx` appends the generated sections after the
front matter, exactly as it does for the pandoc shells.
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# ── Design tokens ────────────────────────────────────────────────────────────
INK = RGBColor(0x1A, 0x1B, 0x2E)
VIOLET = RGBColor(0x7C, 0x5C, 0xD1)
SLATE = RGBColor(0x5B, 0x5B, 0x76)
HAIRLINE = "D8D8E4"
MIST = "F4F3F9"
PAPER = RGBColor(0xFF, 0xFF, 0xFF)

HEAD_FONT = "Segoe UI"
HEAD_FONT_SEMI = "Segoe UI Semibold"
BODY_FONT = "Calibri"

DOC_TYPE_LABELS = {
    "POLICY": "POLICY",
    "STANDARD": "STANDARD",
    "SOP": "STANDARD OPERATING PROCEDURE",
    "PROCEDURE": "PROCEDURE",
    "INCIDENT_RUNBOOK": "INCIDENT RUNBOOK",
    "PLAYBOOK": "PLAYBOOK",
    "PROCESS_FLOW": "PROCESS FLOW",
    "TRAINING": "TRAINING MODULE",
    "TRAINING_MODULE": "TRAINING MODULE",
    "RISK_ASSESSMENT": "RISK ASSESSMENT",
    "AUDIT_PACKAGE": "AUDIT PACKAGE",
    "AI_GOVERNANCE": "AI GOVERNANCE FRAMEWORK",
}


# ── Low-level oxml helpers ───────────────────────────────────────────────────
def _set_char_spacing(run, twentieths: int) -> None:
    """Letter-spacing (tracking), in twentieths of a point."""
    rpr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twentieths))
    rpr.append(spacing)


def _bottom_border(paragraph, color: str, size: int = 6, space: int = 6) -> None:
    """A single hairline rule under a paragraph (used beneath section H1s)."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _shade_cell(cell, color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcpr.append(shd)


def _clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _row_hairline(table, color: str) -> None:
    """Thin horizontal rules between rows only (no verticals)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field (TOC, PAGE, ...) that recomputes in the client."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    run._r.append(begin); run._r.append(instr); run._r.append(sep)
    if placeholder:
        ph = paragraph.add_run(placeholder)
        ph.font.name = BODY_FONT
        ph.font.size = Pt(10)
        ph.font.color.rgb = SLATE
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


# ── Styles ───────────────────────────────────────────────────────────────────
def _configure_styles(doc, accent: RGBColor) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    npf = normal.paragraph_format
    npf.space_after = Pt(6)
    npf.line_spacing = 1.18

    def heading(name, size, color, font=HEAD_FONT_SEMI, before=16, after=4, bold=False):
        st = styles[name]
        st.font.name = font
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        return st

    heading("Heading 1", 14, accent, before=18, after=6)
    heading("Heading 2", 11.5, INK, before=12, after=3)
    heading("Heading 3", 10.5, SLATE, before=10, after=2)

    title = styles["Title"]
    title.font.name = HEAD_FONT
    title.font.size = Pt(30)
    title.font.color.rgb = INK
    title.font.bold = False
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(6)

    def custom(name, size, color, font=HEAD_FONT, bold=False, italic=False):
        try:
            st = styles[name]
        except KeyError:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = font
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.font.italic = italic
        return st

    custom("Eyebrow", 10, accent, font=HEAD_FONT, bold=True)
    custom("Tagline", 10.5, SLATE, font=HEAD_FONT, italic=True)
    custom("CoverMeta", 9.5, SLATE, font=BODY_FONT)
    custom("FrontLabel", 12, accent, font=HEAD_FONT_SEMI, bold=True)


# ── Building blocks ──────────────────────────────────────────────────────────
def _cover(doc, category_label: str, logo_bytes: bytes | None = None) -> None:
    if logo_bytes:
        from io import BytesIO
        try:
            logo_p = doc.add_paragraph()
            logo_p.paragraph_format.space_before = Pt(72)
            logo_p.paragraph_format.space_after = Pt(10)
            logo_p.add_run().add_picture(BytesIO(logo_bytes), width=Inches(1.4))
            top_gap = Pt(6)
        except Exception:
            top_gap = Pt(90)
    else:
        top_gap = Pt(90)

    eyebrow = doc.add_paragraph(style="Eyebrow")
    r = eyebrow.add_run(category_label.upper())
    _set_char_spacing(r, 60)  # 3pt tracking
    eyebrow.paragraph_format.space_before = top_gap
    eyebrow.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph("{{DOCUMENT_TITLE}}", style="Title")
    _bottom_border(title, HAIRLINE, size=6, space=10)
    title.paragraph_format.space_after = Pt(10)

    chip = doc.add_paragraph(style="CoverMeta")
    cr = chip.add_run("CLASSIFICATION:  {{CLASSIFICATION}}")
    _set_char_spacing(cr, 20)
    chip.paragraph_format.space_after = Pt(22)

    # Metadata table (label / value pairs), hairline rows, no box.
    rows = [
        ("Organization", "{{ORGANIZATION_NAME}}"),
        ("Version", "{{VERSION}}"),
        ("Effective Date", "{{EFFECTIVE_DATE}}"),
        ("Next Review", "{{NEXT_REVIEW_DATE}}"),
        ("Owner", "{{POLICY_OWNER}}"),
        ("Approver", "{{APPROVER_NAME}}"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    _row_hairline(table, HAIRLINE)
    for i, (label, value) in enumerate(rows):
        lc, vc = table.rows[i].cells
        lp = lc.paragraphs[0]; lr = lp.add_run(label.upper())
        lr.font.name = HEAD_FONT; lr.font.size = Pt(8.5); lr.font.color.rgb = SLATE
        _set_char_spacing(lr, 20)
        vp = vc.paragraphs[0]; vr = vp.add_run(value)
        vr.font.name = BODY_FONT; vr.font.size = Pt(10.5); vr.font.color.rgb = INK
        lc.width = Inches(1.9); vc.width = Inches(4.1)

    tag = doc.add_paragraph("Audit-ready, never compliant.", style="Tagline")
    tag.paragraph_format.space_before = Pt(26)

    _page_break(doc)


def _contents(doc) -> None:
    label = doc.add_paragraph("Contents", style="FrontLabel")
    label.paragraph_format.space_after = Pt(6)
    toc = doc.add_paragraph()
    _field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and choose “Update Field” to build the contents.")
    _page_break(doc)


def _document_control(doc) -> None:
    label = doc.add_paragraph("Document Control", style="FrontLabel")
    label.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=2, cols=4)
    _row_hairline(table, HAIRLINE)
    headers = ("Version", "Date", "Description of Change", "Author")
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        _shade_cell(c, MIST)
        run = c.paragraphs[0].add_run(h.upper())
        run.font.name = HEAD_FONT; run.font.size = Pt(8.5); run.font.color.rgb = SLATE
        _set_char_spacing(run, 15)
    first = table.rows[1].cells
    for j, val in enumerate(("{{VERSION}}", "{{EFFECTIVE_DATE}}", "Initial issuance.", "{{AUTHOR_NAME}}")):
        run = first[j].paragraphs[0].add_run(val)
        run.font.name = BODY_FONT; run.font.size = Pt(10); run.font.color.rgb = INK
    doc.add_paragraph()


def _page_break(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    run._r.append(br)


def _header_footer(doc, accent: RGBColor) -> None:
    section = doc.sections[0]

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run("{{DOCUMENT_TITLE}}")
    hr.font.name = HEAD_FONT; hr.font.size = Pt(8.5); hr.font.color.rgb = SLATE
    _set_char_spacing(hr, 20)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    left = fp.add_run("{{CLASSIFICATION}}   ")
    left.font.name = HEAD_FONT; left.font.size = Pt(8); left.font.color.rgb = SLATE
    fp.add_run("\t")
    _field(fp, " PAGE ", "")
    fp.add_run(" of ")
    _field(fp, " NUMPAGES ", "")
    for r in fp.runs:
        if not r.font.size:
            r.font.name = BODY_FONT; r.font.size = Pt(8); r.font.color.rgb = SLATE


# ── Public entry point ───────────────────────────────────────────────────────
def build_master_shell(doc_type: str | None, variant: str | None = None, *,
                       accent: RGBColor | None = None, logo_bytes: bytes | None = None):
    """A designed, reusable template shell (front matter + placeholders, no
    sample sections). `_build_docx` fills the placeholders and appends the
    generated sections at Heading 1/2/3."""
    accent = accent or VIOLET
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

    _configure_styles(doc, accent)
    _header_footer(doc, accent)
    _cover(doc, DOC_TYPE_LABELS.get(str(doc_type or "").upper(), "DOCUMENT"), logo_bytes)
    _contents(doc)
    _document_control(doc)

    doc.is_template_shell = True
    doc.is_master_template = True
    doc.section_heading_level = 1  # sections render as styled Heading 1
    return doc
