"""
backend/core/smart_scan_engine.py
Midnight Core — Smart Scan (Bird Eye) Engine

Sits alongside: classifier.py, extractor.py, gap_engine.py, framework_mapper.py
Pipeline:
  1. Template Learning   — parse reference template into section schema
  2. Source Intake       — extract .docx blocks preserving hierarchy
  3. Section ID          — group blocks by semantic heading meaning
  4. Normalization       — strip layout noise, preserve bullets/numbering
  5. LLM Mapping         — Groq semantic section → template alignment
  6. Gap Scoring         — score completeness 0-100, label, flag priorities
"""

import os
import re
import json
import hashlib
from typing import Optional
from enum import Enum

import httpx
from docx import Document
from io import BytesIO


# ---------------------------------------------------------------------------
# Config — reuses existing GROQ env var from your Core pipeline
# ---------------------------------------------------------------------------

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL   = "llama-3.3-70b-versatile"
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")


# ---------------------------------------------------------------------------
# Semantic synonym map — handles messy real-world policy section naming
# ---------------------------------------------------------------------------

SECTION_SYNONYMS: dict[str, list[str]] = {
    "purpose":                ["objective", "intent", "goal", "overview", "introduction", "background"],
    "scope":                  ["applicability", "applies to", "coverage", "audience", "in scope", "out of scope"],
    "policy_statement":       ["statement", "policy intent", "policy position", "policy declaration", "mandate"],
    "definitions":            ["terms", "glossary", "acronyms", "key terms", "definitions and terms"],
    "procedures":             ["process", "steps", "instructions", "guidelines", "how to", "workflow", "process steps"],
    "roles_responsibilities": ["responsibilities", "raci", "roles", "ownership", "accountabilities", "assignments"],
    "references":             ["citations", "related documents", "related policies", "supporting docs", "see also"],
    "revision_history":       ["version history", "change log", "document history", "amendments", "revisions"],
    "exceptions":             ["exception process", "waivers", "deviations", "exemptions"],
    "enforcement":            ["violations", "non-compliance", "consequences", "disciplinary", "sanctions"],
}

SECTION_MIN_WORDS: dict[str, int] = {
    "purpose":                30,
    "scope":                  20,
    "policy_statement":       50,
    "definitions":            15,
    "procedures":             60,
    "roles_responsibilities": 30,
    "references":             10,
    "revision_history":        5,
    "exceptions":             15,
    "enforcement":            15,
}

REQUIRED_SECTIONS = [
    "purpose",
    "scope",
    "policy_statement",
    "procedures",
    "roles_responsibilities",
    "revision_history",
]


# ---------------------------------------------------------------------------
# Result model
# ---------------------------------------------------------------------------

class SectionStatus(str, Enum):
    COMPLETE = "complete"
    PARTIAL  = "partial"
    MISSING  = "missing"


class SmartScanResult:
    def __init__(self):
        self.mapped_sections:   dict  = {}
        self.missing_sections:  list  = []
        self.partial_sections:  list  = []
        self.unmapped_content:  list  = []
        self.section_statuses:  dict  = {}
        self.quality_score:     int   = 0
        self.quality_label:     str   = ""
        self.notes:             list  = []
        self.source_hash:       str   = ""
        self.template_hash:     str   = ""

    def to_dict(self) -> dict:
        return {
            "mapped_sections":   self.mapped_sections,
            "missing_sections":  self.missing_sections,
            "partial_sections":  self.partial_sections,
            "unmapped_content":  self.unmapped_content,
            "section_statuses":  {k: v.value for k, v in self.section_statuses.items()},
            "quality_score":     self.quality_score,
            "quality_label":     self.quality_label,
            "notes":             self.notes,
            "source_hash":       self.source_hash,
            "template_hash":     self.template_hash,
        }


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def _word_count(text: str) -> int:
    return len(text.split())


def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9_]", "_", text.lower().strip())[:40]


def match_section_key(heading: str) -> Optional[str]:
    """Match a heading to a canonical section key via semantic synonyms."""
    h = heading.lower().strip()
    for key, synonyms in SECTION_SYNONYMS.items():
        if key in h or any(s in h for s in synonyms):
            return key
    return None


def extract_docx_blocks(file_bytes: bytes) -> list[dict]:
    """
    Extract all content from a .docx as structured blocks.
    Tables parsed as data (rows/cells), never as layout.
    List hierarchy preserved by style name.
    """
    doc    = Document(BytesIO(file_bytes))
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        level = 0
        if "Heading" in style:
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
        is_list = "List" in style
        blocks.append({
            "type":     "heading" if level > 0 else ("list_item" if is_list else "paragraph"),
            "style":    style,
            "level":    level,
            "text":     text,
            "is_table": False,
            "rows":     [],
        })

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(cells)
        if rows:
            is_kv = all(len(r) <= 2 for r in rows)
            blocks.append({
                "type":     "table_kv" if is_kv else "table_generic",
                "style":    "Table",
                "level":    0,
                "text":     "",
                "is_table": True,
                "rows":     rows,
            })

    return blocks


def learn_template(template_bytes: bytes) -> dict:
    """
    Parse reference template → section schema.
    We OWN the template so layout-aware parsing is safe here.
    """
    doc    = Document(BytesIO(template_bytes))
    schema = {"sections": [], "styles": {}, "hash": _hash_bytes(template_bytes)}

    for para in doc.paragraphs:
        style = para.style.name if para.style else "Normal"
        text  = para.text.strip()
        if not text:
            continue
        if "Heading" in style:
            matched = match_section_key(text)
            schema["sections"].append({
                "id":      matched or _slugify(text),
                "heading": text,
                "matched": matched is not None,
            })

    for para in doc.paragraphs:
        if para.style and para.style.name == "Normal" and para.text.strip():
            if para.runs:
                r = para.runs[0]
                schema["styles"]["font"]      = r.font.name or "Calibri"
                schema["styles"]["font_size"] = str(r.font.size)
            break

    return schema


# ---------------------------------------------------------------------------
# Section identification + normalization
# ---------------------------------------------------------------------------

def identify_sections(blocks: list[dict]) -> dict[str, list[dict]]:
    sections: dict[str, list[dict]] = {}
    current  = "preamble"
    sections[current] = []

    for block in blocks:
        if block["type"] == "heading" and block["level"] <= 2:
            key     = match_section_key(block["text"]) or _slugify(block["text"])
            current = key
            if current not in sections:
                sections[current] = []
        else:
            sections[current].append(block)

    return sections


def normalize_section(blocks: list[dict]) -> str:
    lines = []
    for b in blocks:
        if b["is_table"]:
            for row in b["rows"]:
                lines.append((" | ".join(row)) if len(row) > 2 else (f"{row[0]}: {row[1]}" if len(row) == 2 else row[0]))
        elif b["type"] == "list_item":
            lines.append(f"{'  ' * max(0, b['level']-1)}• {b['text']}")
        else:
            lines.append(b["text"])
    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# LLM semantic mapping
# ---------------------------------------------------------------------------

BIRD_EYE_SYSTEM_PROMPT = """You are Smart Scan (Bird Eye), the semantic analysis engine inside Midnight — Takeoff LLC's enterprise policy intelligence platform.

Your job:
1. Receive normalized text extracted from an unstructured policy document.
2. Receive the list of required template sections.
3. Map the source content to the correct sections based on MEANING, not formatting.
4. Return ONLY valid JSON. No markdown, no preamble, no explanation.

Hard rules:
- NEVER invent or hallucinate content not present in the source.
- NEVER merge unrelated sections.
- Preserve bullet and numbering hierarchy in procedures.
- Use semantic matching: "Objective"→purpose, "Applicability"→scope, "Responsibilities"→roles_responsibilities.
- Mark missing sections as missing. Do not fill gaps.
- Mark partially complete sections as partial.

Output (strict JSON, no fences):
{
  "mapped_sections": {
    "purpose": "...",
    "scope": "...",
    "policy_statement": "...",
    "procedures": "...",
    "roles_responsibilities": "...",
    "definitions": "...",
    "references": "...",
    "revision_history": "...",
    "exceptions": "...",
    "enforcement": "..."
  },
  "missing_sections": ["section_key"],
  "partial_sections": ["section_key"],
  "unmapped_content": ["text not matching any section"],
  "notes": ["observation about document quality or structure"]
}
Only include keys in mapped_sections where content was actually found."""


async def llm_map_sections(extracted_text: str, template_sections: list[str]) -> dict:
    user_msg = (
        f"TEMPLATE SECTIONS REQUIRED:\n{json.dumps(template_sections, indent=2)}\n\n"
        f"SOURCE DOCUMENT CONTENT:\n---\n{extracted_text[:12000]}\n---\n\n"
        f"Map the source content to the template sections. Return only JSON."
    )
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            GROQ_API_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={
                "model":       GROQ_MODEL,
                "temperature": 0.1,
                "max_tokens":  4096,
                "messages": [
                    {"role": "system", "content": BIRD_EYE_SYSTEM_PROMPT},
                    {"role": "user",   "content": user_msg},
                ],
            },
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()

    raw = re.sub(r"^```[a-z]*\n?", "", raw)
    raw = re.sub(r"\n?```$",       "", raw)
    return json.loads(raw)


# ---------------------------------------------------------------------------
# Gap scoring
# ---------------------------------------------------------------------------

def score_document(
    mapped:   dict,
    missing:  list,
    partial:  list,
) -> tuple[int, str, dict]:
    statuses: dict[str, SectionStatus] = {}
    total = earned = 0

    for key in REQUIRED_SECTIONS:
        w      = 15 if key in ("policy_statement", "procedures") else 10
        total += w
        if key in missing:
            statuses[key] = SectionStatus.MISSING
        elif key in partial:
            statuses[key] = SectionStatus.PARTIAL
            earned       += w * 0.5
        elif key in mapped:
            wc = _word_count(mapped[key] if isinstance(mapped[key], str) else str(mapped[key]))
            if wc >= SECTION_MIN_WORDS.get(key, 20):
                statuses[key] = SectionStatus.COMPLETE
                earned       += w
            else:
                statuses[key] = SectionStatus.PARTIAL
                earned       += w * 0.4
        else:
            statuses[key] = SectionStatus.MISSING

    score = int((earned / total) * 100) if total else 0
    label = "Strong" if score >= 80 else "Moderate" if score >= 55 else "Weak" if score >= 30 else "Critical"
    return score, label, statuses


def build_notes(statuses: dict, mapped: dict, missing: list, partial: list) -> list[str]:
    notes = []
    for key in missing:
        if key in REQUIRED_SECTIONS:
            notes.append(f"REQUIRED — '{key.replace('_',' ').title()}' section not found in source document.")
    for key in partial:
        notes.append(f"PARTIAL — '{key.replace('_',' ').title()}' present but lacks sufficient detail.")
    if "procedures" in mapped:
        c = str(mapped["procedures"])
        if "step" not in c.lower() and "•" not in c and "1." not in c:
            notes.append("Procedures section may lack step-by-step structure — consider explicit numbered steps.")
    if "revision_history" in missing:
        notes.append("No revision history found — document cannot be version-tracked.")
    if "roles_responsibilities" in missing:
        notes.append("No roles or RACI defined — policy ownership is unclear.")
    return notes


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

async def run_smart_scan(
    source_bytes:    bytes,
    template_bytes:  bytes,
    source_filename: str = "source.docx",
) -> SmartScanResult:
    result               = SmartScanResult()
    result.source_hash   = _hash_bytes(source_bytes)
    result.template_hash = _hash_bytes(template_bytes)

    # 1. Learn template
    schema            = learn_template(template_bytes)
    template_sections = [s["id"] for s in schema["sections"]] or list(SECTION_SYNONYMS.keys())

    # 2. Extract + identify sections from source
    blocks      = extract_docx_blocks(source_bytes)
    raw_sections = identify_sections(blocks)

    # 3. Normalize to text
    parts = []
    for key, blks in raw_sections.items():
        content = normalize_section(blks)
        if content:
            parts.append(f"[SECTION: {key}]\n{content}")
    full_text = "\n\n".join(parts)

    # 4. LLM mapping
    llm = await llm_map_sections(full_text, template_sections)

    result.mapped_sections  = llm.get("mapped_sections", {})
    result.missing_sections = llm.get("missing_sections", [])
    result.partial_sections = llm.get("partial_sections", [])
    result.unmapped_content = llm.get("unmapped_content", [])

    # 5. Score + notes
    score, label, statuses = score_document(
        result.mapped_sections,
        result.missing_sections,
        result.partial_sections,
    )
    result.quality_score    = score
    result.quality_label    = label
    result.section_statuses = statuses
    result.notes            = llm.get("notes", []) + build_notes(
        statuses, result.mapped_sections, result.missing_sections, result.partial_sections
    )

    return result
