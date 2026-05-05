"""
Midnight Core - Pipeline Routes
Takeoff LLC
"""

import json
import os
import base64
import csv
import logging
import mimetypes
import re
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, StreamingResponse
from pydantic import BaseModel

from backend.core.framework_layer import (
    build_framework_mapping_rules,
    build_framework_prompt_context,
)
from backend.core.json_parser import (
    ParsedModelOutputError,
    PolicySchemaError,
    normalize_policy_payload,
    parse_model_json,
)
from backend.renderers.pdf_renderer import build_grc_summary_pdf
from backend.storage.file_store import (
    SupabaseStoreError,
    count_activity_for_tenant,
    create_activity_event,
    create_onboarding_session,
    create_tenant,
    get_onboarding_session,
    get_tenant,
    get_tenant_by_slug,
    list_generated_documents,
    replace_enabled_modules,
    save_policy_draft,
    save_generated_document,
    update_onboarding_session,
    update_profile_membership,
)

try:
    import anthropic
except ImportError:  # pragma: no cover
    anthropic = None

try:
    import requests as _requests
except ImportError:  # pragma: no cover
    _requests = None

load_dotenv()

logger = logging.getLogger("midnight.policy_json")

router = APIRouter()
pipeline_router = APIRouter(prefix="/pipeline", tags=["pipeline"])
org_router = APIRouter(prefix="/api/orgs", tags=["orgs"])

GO_SERVICE_URL = os.getenv("GO_SERVICE_URL", "http://localhost:8080")

_FW_TO_GO = {
    "HIPAA": "HIPAA",
    "PCI DSS": "PCI_DSS",
    "NIST CSF": "NIST_CSF",
    "SOC 2": "SOC2",
    "HITRUST": "HITRUST",
    "ISO 27001": "ISO_27001",
}

TRIAL_MAX_UPLOADS = 3
TRIAL_MAX_FRAMEWORKS = 1
TRIAL_WATERMARK_TEXT = "TRIAL - Midnight Preview"
POLICY_SCHEMA_VERSION = "midnight-policy-v2"
POLICY_REQUIRED_SLOTS = [
    "purpose",
    "scope",
    "definitions",
    "roles_responsibilities",
    "policy_statement",
    "procedures",
    "compliance_requirements",
    "exceptions",
    "review_cycle",
    "approval",
]
POLICY_SLOT_SPECS = [
    {
        "slot_id": "purpose",
        "heading": "Purpose",
        "instruction": "State why the policy exists, the business objective, and the risk or compliance need it addresses.",
    },
    {
        "slot_id": "scope",
        "heading": "Scope",
        "instruction": "Define which people, systems, data, locations, and business processes the policy applies to.",
    },
    {
        "slot_id": "definitions",
        "heading": "Definitions",
        "instruction": "List and explain the specific terms, acronyms, or concepts needed to understand the policy.",
    },
    {
        "slot_id": "roles_responsibilities",
        "heading": "Roles and Responsibilities",
        "instruction": "Identify the accountable roles and what each role is responsible for under this policy.",
    },
    {
        "slot_id": "policy_statement",
        "heading": "Policy Statement",
        "instruction": "Write the core mandatory policy requirements in direct compliance-ready language.",
    },
    {
        "slot_id": "procedures",
        "heading": "Procedures",
        "instruction": "Describe the operational steps, required actions, or workflow expectations needed to carry out the policy.",
    },
    {
        "slot_id": "compliance_requirements",
        "heading": "Compliance Requirements",
        "instruction": "State the control obligations, evidence expectations, and framework-related requirements that support compliance.",
    },
    {
        "slot_id": "exceptions",
        "heading": "Exceptions",
        "instruction": "Explain how exceptions are requested, approved, documented, and reviewed.",
    },
    {
        "slot_id": "review_cycle",
        "heading": "Review Cycle",
        "instruction": "Describe how often this policy must be reviewed, updated, and re-approved.",
    },
    {
        "slot_id": "approval",
        "heading": "Approval",
        "instruction": "Describe the approval authority, approval workflow, and what constitutes policy approval.",
    },
]
SECTION_CONTENT_LIMIT = 12000


def _slugify_value(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", str(value or "").strip().lower()).strip("-")
    return slug or "tenant"


def _tenant_id_from_request(request: Request) -> str:
    tenant_id = getattr(request.state, "tenant_id", None)
    if not tenant_id:
        raise HTTPException(status_code=401, detail="Authenticated tenant context is missing.")
    return tenant_id


def _tenant_context_from_request(request: Request) -> dict:
    auth_context = getattr(request.state, "auth_context", None) or {}
    tenant_id = auth_context.get("tenant_id") or getattr(request.state, "tenant_id", None)
    if not tenant_id:
        raise HTTPException(status_code=401, detail="Authenticated tenant context is missing.")
    try:
        tenant = get_tenant(tenant_id)
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    if tenant is None:
        raise HTTPException(status_code=403, detail="Tenant record not found.")
    return tenant


def _assert_tenant_access(request: Request, tenant_id: str) -> dict:
    current_tenant_id = _tenant_id_from_request(request)
    if str(current_tenant_id) != str(tenant_id):
        raise HTTPException(status_code=403, detail="Cross-tenant access is not allowed.")
    return _tenant_context_from_request(request)


def _enforce_trial_limits(
    request: Request,
    *,
    frameworks: list[str] | None = None,
    upload_action: str | None = None,
) -> dict:
    tenant = _tenant_context_from_request(request)
    if str(tenant.get("plan_type") or "").lower() != "trial":
        return tenant

    if frameworks is not None and len([fw for fw in frameworks if fw]) > TRIAL_MAX_FRAMEWORKS:
        raise HTTPException(
            status_code=403,
            detail=f"Trial plans support only {TRIAL_MAX_FRAMEWORKS} framework at a time. Upgrade to continue.",
        )

    if upload_action:
        try:
            upload_count = count_activity_for_tenant(tenant["id"], action=upload_action)
        except SupabaseStoreError as exc:
            raise HTTPException(status_code=503, detail=str(exc)) from exc
        if upload_count >= TRIAL_MAX_UPLOADS:
            raise HTTPException(
                status_code=403,
                detail=f"Trial plans are limited to {TRIAL_MAX_UPLOADS} uploads. Upgrade to continue.",
            )
    return tenant


def _watermark_docx_bytes(docx_bytes: bytes, *, watermark_text: str = TRIAL_WATERMARK_TEXT) -> bytes:
    doc = Document(BytesIO(docx_bytes))
    notice = doc.paragraphs[0].insert_paragraph_before()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run(watermark_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    run.font.size = Pt(14)

    footer = doc.sections[0].footer.paragraphs[0]
    existing_footer = footer.text.strip()
    footer.text = f"{watermark_text} · {existing_footer}" if existing_footer else watermark_text
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _go_framework_coverage(document: str, frameworks: list[str], title: str = "") -> list | None:
    if _requests is None:
        return None
    go_frameworks = [_FW_TO_GO[f] for f in frameworks if f in _FW_TO_GO]
    if not go_frameworks:
        return None
    try:
        resp = _requests.post(
            f"{GO_SERVICE_URL}/analyze",
            json={"title": title, "document": document, "frameworks": go_frameworks},
            timeout=180,
        )
        resp.raise_for_status()
        return resp.json().get("frameworks", [])
    except Exception:
        return None


@pipeline_router.get("/smoke-docx")
async def smoke_docx():
    path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.docx")

    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(path)

    return FileResponse(
        path=path,
        filename="smoke_test.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
ANTHROPIC_MODEL = "claude-opus-4-5"
SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md"}
SUPPORTED_TEMPLATE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
    ".txt",
    ".md",
}

MIGRATE_OUTPUT_SECTION_LABELS = {
    "purpose": "Purpose and Scope",
    "definitions": "Definitions",
    "policy_statement": "Policy Statement",
    "procedures": "Procedures",
    "related_policies": "Related Policies",
    "citations_references": "Citations / References",
    "revision_history": "Revision History",
}

MIGRATE_SECTION_ALIASES = {
    "purpose": ["Purpose", "Purpose and Scope", "Background", "Overview", "Objective", "Introduction"],
    "definitions": ["Glossary", "Key Terms", "Definitions", "Terminology"],
    "policy_statement": [
        "Policy Statement",
        "Policy Body",
        "Policy Details",
        "Policy Content",
        "Policy Requirements",
        "Rules",
    ],
    "procedures": ["Procedures", "Process Steps", "Controls", "Implementation", "Guidelines"],
    "related_policies": ["Related Policies", "See Also", "References", "Associated Documents"],
    "citations_references": [
        "Citations / References",
        "Standards",
        "Regulatory References",
        "Citations",
        "Compliance References",
    ],
    "revision_history": ["Revision History", "Change Log", "Document History", "Version History"],
}

MIGRATE_HEADER_ALIASES = {
    "policy_name": ["Policy Name"],
    "policy_number": ["Policy Number"],
    "version": ["Version"],
    "grc_id": ["GRC ID"],
    "effective_date": ["Effective Date"],
    "supersedes_policy": ["Supersedes Policy"],
    "last_reviewed_date": ["Last Reviewed Date", "Last Reviewed"],
    "last_revised_date": ["Last Revised Date", "Last Revised"],
    "policy_custodian": ["Policy Custodian"],
    "policy_owner": ["Policy Owner"],
    "policy_approver": ["Policy Approver"],
}


def _get_workspace_id(request: Request) -> str:
    return _tenant_id_from_request(request)


def _get_anthropic_client():
    if anthropic is None:
        raise HTTPException(
            status_code=503,
            detail="Anthropic dependency is not installed on the server.",
        )
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY is not configured.")
    return anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


def _require_supported_file(upload: UploadFile, field_name: str = "file") -> None:
    filename = (upload.filename or field_name).lower()
    ext = "." + filename.rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"{field_name} must be .docx, .txt, or .md. Got '{upload.filename}'.",
        )


def _require_supported_template_file(upload: UploadFile, field_name: str = "template_file") -> None:
    filename = (upload.filename or field_name).lower()
    ext = "." + filename.rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in SUPPORTED_TEMPLATE_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"{field_name} must be one of "
                ".png, .jpg, .jpeg, .pdf, .docx, .xlsx, .csv, .txt, or .md. "
                f"Got '{upload.filename}'."
            ),
        )


def _extract_text_from_upload(file_bytes: bytes, filename: str) -> str:
    lower = filename.lower()

    if lower.endswith(".docx"):
        doc = Document(BytesIO(file_bytes))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        return "\n".join(lines)

    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _iter_docx_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield "table", Table(child, doc)


def _table_rows_to_lists(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        values = [cell.text or "" for cell in row.cells]
        if any(value.strip() for value in values):
            rows.append(values)
    return rows


def _extract_source_blocks(file_bytes: bytes, filename: str) -> list[dict]:
    lower = filename.lower()
    blocks: list[dict] = []

    if lower.endswith(".docx"):
        doc = Document(BytesIO(file_bytes))
        for kind, item in _iter_docx_blocks(doc):
            if kind == "paragraph":
                blocks.append({"kind": "paragraph", "text": item.text or ""})
            else:
                rows = _table_rows_to_lists(item)
                if rows:
                    blocks.append(
                        {
                            "kind": "table",
                            "rows": rows,
                            "text": "\n".join(" | ".join(row) for row in rows),
                        }
                    )
        return blocks

    text = _extract_text_from_upload(file_bytes, filename)
    for line in text.splitlines():
        blocks.append({"kind": "paragraph", "text": line})
    return blocks


def _normalize_source_label(text: str) -> str:
    normalized = re.sub(r"^[#>\-\*\s]+", "", str(text or "").strip())
    normalized = normalized.rstrip(":").strip().lower()
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized


def _split_semicolons(value: str) -> str:
    return str(value or "").replace(";", "\n")


def _match_header_value(text: str) -> tuple[Optional[str], Optional[str]]:
    raw = str(text or "").strip()
    if not raw:
        return None, None

    for output_key, aliases in MIGRATE_HEADER_ALIASES.items():
        for alias in aliases:
            match = re.match(
                rf"^\s*{re.escape(alias)}\s*[:\-–]\s*(.+?)\s*$",
                raw,
                flags=re.IGNORECASE,
            )
            if match:
                return output_key, match.group(1)
    return None, None


def _match_section_heading(text: str) -> tuple[Optional[str], Optional[str]]:
    raw = str(text or "").strip()
    if not raw:
        return None, None

    normalized = _normalize_source_label(raw)
    for output_key, aliases in MIGRATE_SECTION_ALIASES.items():
        for alias in aliases:
            alias_normalized = _normalize_source_label(alias)
            if normalized == alias_normalized:
                return output_key, None

            match = re.match(
                rf"^\s*{re.escape(alias)}\s*[:\-–]\s*(.+?)\s*$",
                raw,
                flags=re.IGNORECASE,
            )
            if match:
                return output_key, match.group(1)
    return None, None


def _capture_header_rows(rows: list[list[str]], header_values: dict) -> bool:
    captured = False
    for row in rows:
        if len(row) < 2:
            continue
        key, value = _match_header_value(f"{row[0]}: {row[1]}")
        if key and value and not header_values.get(key):
            header_values[key] = value
            captured = True
    return captured


def _looks_like_revision_history_header(row: list[str]) -> bool:
    normalized = [_normalize_source_label(cell) for cell in row if str(cell or "").strip()]
    if not normalized:
        return False
    header_tokens = {"version", "date", "description", "change", "change log", "history"}
    return sum(token in header_tokens for token in normalized) >= 2


def _build_revision_history(rows: list[list[str]]) -> list[dict]:
    entries: list[dict] = []
    filtered_rows = [row for row in rows if any(str(cell or "").strip() for cell in row)]
    if filtered_rows and _looks_like_revision_history_header(filtered_rows[0]):
        filtered_rows = filtered_rows[1:]

    for row in filtered_rows:
        cleaned = [_split_semicolons(cell).strip() for cell in row]
        if not any(cleaned):
            continue
        entries.append(
            {
                "version": cleaned[0] if len(cleaned) > 0 else "",
                "date": cleaned[1] if len(cleaned) > 1 else "",
                "description": " | ".join(cell for cell in cleaned[2:] if cell) if len(cleaned) > 2 else "",
            }
        )
    return entries


def _join_verbatim_content(parts: list[str]) -> str:
    if not parts:
        return ""
    text = "\n".join(parts)
    return text.strip("\n")


def _map_migrate_source(file_bytes: bytes, filename: str) -> dict:
    blocks = _extract_source_blocks(file_bytes, filename)
    header_values: dict[str, str] = {}
    section_parts: dict[str, list[str]] = {key: [] for key in MIGRATE_SECTION_ALIASES}
    revision_rows: list[list[str]] = []
    current_section: Optional[str] = None

    for block in blocks:
        if block["kind"] == "paragraph":
            text = str(block.get("text", ""))
            if not text.strip():
                if current_section and current_section != "revision_history" and section_parts[current_section]:
                    section_parts[current_section].append("")
                continue

            header_key, header_value = _match_header_value(text)
            if header_key and header_value and not header_values.get(header_key):
                header_values[header_key] = header_value
                continue

            section_key, inline_value = _match_section_heading(text)
            if section_key:
                current_section = section_key
                if inline_value:
                    if section_key == "revision_history":
                        revision_rows.append([_split_semicolons(inline_value)])
                    else:
                        section_parts[section_key].append(_split_semicolons(inline_value))
                continue

            if current_section == "revision_history":
                revision_rows.append([_split_semicolons(text)])
            elif current_section:
                section_parts[current_section].append(_split_semicolons(text))
            continue

        rows = block.get("rows", [])
        if not rows:
            continue

        if _capture_header_rows(rows, header_values) and current_section is None:
            continue

        if current_section == "revision_history":
            revision_rows.extend(rows)
        elif current_section:
            for row in rows:
                section_parts[current_section].append(_split_semicolons(" | ".join(row)))

    related_policies = _join_verbatim_content(section_parts["related_policies"])
    citations_references = _join_verbatim_content(section_parts["citations_references"])
    combined_references_parts = []
    if related_policies:
        combined_references_parts.append(f"Related Policies\n{related_policies}")
    if citations_references:
        combined_references_parts.append(f"Citations / References\n{citations_references}")

    return {
        "headers": header_values,
        "purpose": _join_verbatim_content(section_parts["purpose"]),
        "definitions": _join_verbatim_content(section_parts["definitions"]),
        "policy_statement": _join_verbatim_content(section_parts["policy_statement"]),
        "procedures": _join_verbatim_content(section_parts["procedures"]),
        "related_policies": related_policies,
        "citations_references": citations_references,
        "combined_references": "\n\n".join(part for part in combined_references_parts if part),
        "revision_history": _build_revision_history(revision_rows),
    }


def _build_migrate_mapping_context(mapped_source: dict) -> str:
    header_lines = []
    for key in (
        "policy_name",
        "policy_number",
        "version",
        "grc_id",
        "effective_date",
        "supersedes_policy",
        "last_reviewed_date",
        "last_revised_date",
        "policy_custodian",
        "policy_owner",
        "policy_approver",
    ):
        value = mapped_source.get("headers", {}).get(key)
        if value:
            header_lines.append(f"{key}: {value}")

    section_lines = []
    for key in (
        "purpose",
        "definitions",
        "policy_statement",
        "procedures",
        "related_policies",
        "citations_references",
    ):
        value = mapped_source.get(key)
        if value:
            section_lines.append(f"{key.upper()}:\n{value}")

    revision_history = mapped_source.get("revision_history") or []
    if revision_history:
        revision_lines = ["REVISION_HISTORY:"]
        for entry in revision_history:
            revision_lines.append(
                f"{entry.get('version', '')} | {entry.get('date', '')} | {entry.get('description', '')}"
            )
        section_lines.append("\n".join(revision_lines))

    return "\n\n".join(
        part for part in ["\n".join(header_lines).strip(), "\n\n".join(section_lines).strip()] if part
    )


def _apply_migrate_source_mapping(policy_data: dict, mapped_source: dict) -> dict:
    updated = dict(policy_data or {})
    headers = mapped_source.get("headers", {}) or {}

    if headers.get("policy_name"):
        updated["policy_name"] = headers["policy_name"]
    if headers.get("version"):
        updated["version"] = headers["version"]
    if headers.get("effective_date"):
        updated["effective_date"] = headers["effective_date"]
    if headers.get("policy_owner"):
        updated["owner"] = headers["policy_owner"]

    for field in (
        "policy_number",
        "grc_id",
        "supersedes_policy",
        "last_reviewed_date",
        "last_revised_date",
        "policy_custodian",
        "policy_owner",
        "policy_approver",
    ):
        if headers.get(field):
            updated[field] = headers[field]

    if mapped_source.get("purpose"):
        updated["purpose"] = mapped_source["purpose"]
    if mapped_source.get("definitions"):
        updated["definitions"] = mapped_source["definitions"]
    if mapped_source.get("policy_statement"):
        updated["policy_statement"] = mapped_source["policy_statement"]
    if mapped_source.get("procedures"):
        updated["procedures"] = mapped_source["procedures"]
    if mapped_source.get("related_policies"):
        updated["related_policies"] = mapped_source["related_policies"]
    if mapped_source.get("citations_references"):
        updated["citations_references"] = mapped_source["citations_references"]
    if mapped_source.get("combined_references"):
        updated["references"] = mapped_source["combined_references"]
    if mapped_source.get("revision_history"):
        updated["revision_history"] = mapped_source["revision_history"]

    return updated


def _guess_media_type(filename: str, default: str = "application/octet-stream") -> str:
    return mimetypes.guess_type(filename or "")[0] or default


def _canonical_template_section_key(label: str) -> Optional[str]:
    normalized = _normalize_source_label(label)
    if not normalized:
        return None

    for key, output_label in MIGRATE_OUTPUT_SECTION_LABELS.items():
        if normalized == _normalize_source_label(output_label):
            return key

    for key, aliases in MIGRATE_SECTION_ALIASES.items():
        for alias in aliases:
            if normalized == _normalize_source_label(alias):
                return key
    return None


def _empty_template_section_map(template_name: str, template_type: str, mode: str) -> dict:
    return {
        "template_name": template_name,
        "template_type": template_type,
        "extraction_mode": mode,
        "header_fields": [],
        "sections": [],
        "field_mappings": [],
        "section_order": [],
        "raw_outline": [],
    }


def _finalize_template_section_map(section_map: dict) -> dict:
    ordered_sections: list[str] = []
    for section in section_map.get("sections", []):
        label = section.get("output_section")
        if label and label not in ordered_sections:
            ordered_sections.append(label)
    section_map["section_order"] = ordered_sections
    return section_map


def _append_template_heading(section_map: dict, label: str, source_type: str, *, structure: str = "heading"):
    canonical_key = _canonical_template_section_key(label)
    output_label = MIGRATE_OUTPUT_SECTION_LABELS.get(canonical_key) if canonical_key else None
    entry = {
        "detected_label": label.strip(),
        "output_section": output_label or label.strip(),
        "canonical_key": canonical_key,
        "source_type": source_type,
        "structure": structure,
    }
    section_map["sections"].append(entry)
    section_map["raw_outline"].append(label.strip())


def _append_template_header_field(section_map: dict, label: str, sample_value: str = ""):
    canonical_key, _ = _match_header_value(f"{label}: {sample_value or 'value'}")
    entry = {
        "detected_label": label.strip(),
        "canonical_key": canonical_key,
        "sample_value": sample_value.strip(),
    }
    section_map["header_fields"].append(entry)


def _extract_markdown_text_section_map(text: str, filename: str, template_type: str) -> dict:
    section_map = _empty_template_section_map(filename, template_type, "text-parse")
    lines = text.splitlines()

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()
        if not stripped:
            continue

        header_key, header_value = _match_header_value(stripped)
        if header_key:
            _append_template_header_field(section_map, stripped.split(":", 1)[0], header_value or "")
            continue

        heading_label = None
        if stripped.startswith("#"):
            heading_label = stripped.lstrip("#").strip()
        else:
            section_key, _ = _match_section_heading(stripped)
            if section_key:
                heading_label = stripped.rstrip(":")
            elif len(stripped) < 80 and re.fullmatch(r"[A-Za-z0-9 /&\-\(\)]+:?", stripped):
                if _canonical_template_section_key(stripped):
                    heading_label = stripped.rstrip(":")

        if heading_label:
            _append_template_heading(section_map, heading_label, template_type)

    return _finalize_template_section_map(section_map)


def _extract_docx_template_section_map(file_bytes: bytes, filename: str) -> dict:
    section_map = _empty_template_section_map(filename, "docx", "docx-parse")
    doc = Document(BytesIO(file_bytes))

    for kind, item in _iter_docx_blocks(doc):
        if kind == "paragraph":
            text = (item.text or "").strip()
            if not text:
                continue

            header_key, header_value = _match_header_value(text)
            if header_key:
                _append_template_header_field(section_map, text.split(":", 1)[0], header_value or "")
                continue

            style_name = getattr(getattr(item, "style", None), "name", "") or ""
            if "Heading" in style_name or _canonical_template_section_key(text):
                _append_template_heading(section_map, text.rstrip(":"), "docx")
        else:
            rows = _table_rows_to_lists(item)
            if not rows:
                continue
            first_row = rows[0]
            if _capture_header_rows(rows, {}):
                for row in rows:
                    if len(row) >= 2:
                        key, value = _match_header_value(f"{row[0]}: {row[1]}")
                        if key:
                            _append_template_header_field(section_map, row[0], row[1])
                continue

            for row in rows:
                label = next((str(cell).strip() for cell in row if str(cell).strip()), "")
                if not label:
                    continue
                canonical = _canonical_template_section_key(label)
                if canonical:
                    _append_template_heading(section_map, label.rstrip(":"), "docx", structure="table")

            if first_row:
                section_map["field_mappings"].append(
                    {
                        "columns": [str(cell).strip() for cell in first_row],
                        "rows": rows[1:] if len(rows) > 1 else [],
                    }
                )

    return _finalize_template_section_map(section_map)


def _extract_csv_template_section_map(file_bytes: bytes, filename: str) -> dict:
    text = _extract_text_from_upload(file_bytes, filename)
    reader = csv.reader(text.splitlines())
    rows = [row for row in reader if any(str(cell).strip() for cell in row)]
    section_map = _empty_template_section_map(filename, "csv", "tabular-parse")
    if not rows:
        return section_map

    header_row = [str(cell).strip() for cell in rows[0]]
    section_map["field_mappings"].append({"columns": header_row, "rows": rows[1:]})

    for row in rows[1:] if len(rows) > 1 else []:
        label = next((str(cell).strip() for cell in row if str(cell).strip()), "")
        if not label:
            continue
        if _canonical_template_section_key(label):
            _append_template_heading(section_map, label.rstrip(":"), "csv", structure="row")

    return _finalize_template_section_map(section_map)


def _xlsx_shared_strings(zf: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    values = []
    for si in root.findall("a:si", ns):
        text = "".join(node.text or "" for node in si.findall(".//a:t", ns))
        values.append(text)
    return values


def _xlsx_rows(file_bytes: bytes) -> list[list[str]]:
    with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
        shared_strings = _xlsx_shared_strings(zf)
        worksheet_names = sorted(name for name in zf.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"))
        if not worksheet_names:
            return []

        root = ET.fromstring(zf.read(worksheet_names[0]))
        ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        rows: list[list[str]] = []
        for row in root.findall(".//a:sheetData/a:row", ns):
            values: list[str] = []
            for cell in row.findall("a:c", ns):
                cell_type = cell.attrib.get("t")
                value_node = cell.find("a:v", ns)
                if value_node is None:
                    values.append("")
                    continue
                raw = value_node.text or ""
                if cell_type == "s":
                    try:
                        values.append(shared_strings[int(raw)])
                    except Exception:
                        values.append(raw)
                else:
                    values.append(raw)
            if any(str(value).strip() for value in values):
                rows.append(values)
        return rows


def _extract_xlsx_template_section_map(file_bytes: bytes, filename: str) -> dict:
    rows = _xlsx_rows(file_bytes)
    section_map = _empty_template_section_map(filename, "xlsx", "tabular-parse")
    if not rows:
        return section_map

    header_row = [str(cell).strip() for cell in rows[0]]
    section_map["field_mappings"].append({"columns": header_row, "rows": rows[1:]})

    for row in rows[1:] if len(rows) > 1 else []:
        label = next((str(cell).strip() for cell in row if str(cell).strip()), "")
        if label and _canonical_template_section_key(label):
            _append_template_heading(section_map, label.rstrip(":"), "xlsx", structure="row")

    return _finalize_template_section_map(section_map)


def _extract_pdf_text_basic(file_bytes: bytes) -> str:
    text_chunks = []
    for match in re.findall(rb"\((.*?)\)", file_bytes, flags=re.DOTALL):
        decoded = match.replace(rb"\(", b"(").replace(rb"\)", b")").replace(rb"\\n", b"\n")
        candidate = decoded.decode("latin-1", errors="ignore")
        if any(ch.isalpha() for ch in candidate):
            text_chunks.append(candidate)
    return "\n".join(text_chunks)


TEMPLATE_MAP_SYSTEM_PROMPT = """You are a template structure extraction engine for Midnight.

Return valid JSON only with this structure:
{
  "template_name": "...",
  "template_type": "...",
  "extraction_mode": "...",
  "header_fields": [
    {"detected_label": "...", "canonical_key": "...", "sample_value": "..."}
  ],
  "sections": [
    {
      "detected_label": "...",
      "output_section": "...",
      "canonical_key": "...",
      "source_type": "...",
      "structure": "heading|table|field_map|image"
    }
  ],
  "field_mappings": [],
  "section_order": ["..."],
  "raw_outline": ["..."]
}

Recognize these canonical output sections:
- Purpose and Scope
- Definitions
- Policy Statement
- Procedures
- Related Policies
- Citations / References
- Revision History

Recognize these header fields wherever present:
- Policy Name
- Policy Number
- Version
- GRC ID
- Effective Date
- Supersedes Policy
- Last Reviewed Date
- Last Revised Date
- Policy Custodian
- Policy Owner
- Policy Approver

Rules:
- Identify visible section headers and layout structure only.
- Do not invent sections that are not present.
- Return JSON only, no markdown."""


def _parse_template_map_json(raw_text: str) -> dict:
    data = _parse_generic_model_object(raw_text, flow="template_section_map")
    return _finalize_template_section_map(data)


def _extract_template_section_map_with_model(
    *,
    template_bytes: bytes,
    filename: str,
    source_type: str,
) -> dict:
    client = _get_anthropic_client()
    media_type = _guess_media_type(filename, "application/octet-stream")
    encoded = base64.b64encode(template_bytes).decode("utf-8")

    if source_type in {"png", "jpg", "jpeg"}:
        content = [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": encoded,
                },
            },
            {
                "type": "text",
                "text": f"Extract the section map for template file '{filename}'.",
            },
        ]
    else:
        content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": encoded,
                },
            },
            {
                "type": "text",
                "text": f"Extract the section map for template file '{filename}'.",
            },
        ]

    message = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=2048,
        system=TEMPLATE_MAP_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
    )
    return _parse_template_map_json(message.content[0].text)


def _default_template_section_map(template_name: str) -> dict:
    section_map = _empty_template_section_map(template_name, "named-template", "default")
    for key in (
        "purpose",
        "definitions",
        "policy_statement",
        "procedures",
        "related_policies",
        "citations_references",
        "revision_history",
    ):
        _append_template_heading(section_map, MIGRATE_OUTPUT_SECTION_LABELS[key], "default")
    return _finalize_template_section_map(section_map)


def _extract_template_section_map(
    *,
    template_name: str,
    template_file_bytes: Optional[bytes] = None,
    template_filename: Optional[str] = None,
) -> dict:
    if not template_file_bytes or not template_filename:
        return _default_template_section_map(template_name)

    lower = template_filename.lower()
    ext = "." + lower.rsplit(".", 1)[-1] if "." in lower else ""
    source_type = ext.lstrip(".")

    if ext == ".docx":
        return _extract_docx_template_section_map(template_file_bytes, template_filename)
    if ext == ".csv":
        return _extract_csv_template_section_map(template_file_bytes, template_filename)
    if ext == ".xlsx":
        return _extract_xlsx_template_section_map(template_file_bytes, template_filename)
    if ext in {".txt", ".md"}:
        text = _extract_text_from_upload(template_file_bytes, template_filename)
        return _extract_markdown_text_section_map(text, template_filename, source_type)
    if ext == ".pdf":
        text = _extract_pdf_text_basic(template_file_bytes)
        parsed = _extract_markdown_text_section_map(text, template_filename, "pdf")
        if parsed.get("sections") or parsed.get("header_fields"):
            return parsed
        return _extract_template_section_map_with_model(
            template_bytes=template_file_bytes,
            filename=template_filename,
            source_type="pdf",
        )
    if ext in {".png", ".jpg", ".jpeg"}:
        return _extract_template_section_map_with_model(
            template_bytes=template_file_bytes,
            filename=template_filename,
            source_type=source_type,
        )

    return _default_template_section_map(template_name)


def _strip_json_fences(raw: str) -> str:
    return raw.replace("```json", "").replace("```", "").strip()


def _log_model_output_failure(
    *,
    flow: str,
    raw_text: str,
    error: Exception,
    context: dict[str, object] | None = None,
) -> None:
    payload = {
        "flow": flow,
        "error": str(error),
        "raw_output_excerpt": str(raw_text or "")[:4000],
    }
    if context:
        payload.update(context)
    logger.warning("model_json_parse_failed", extra=payload)


def _parse_policy_model_output(
    raw_text: str,
    *,
    flow: str,
    organization_hint: str,
    required_frameworks: list[str] | None = None,
    context: dict[str, object] | None = None,
) -> dict:
    try:
        parsed = parse_model_json(raw_text)
    except ParsedModelOutputError as exc:
        _log_model_output_failure(flow=flow, raw_text=raw_text, error=exc, context=context)
        raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {exc}") from exc

    try:
        return normalize_policy_payload(
            parsed,
            organization_hint=organization_hint,
            required_frameworks=required_frameworks,
        )
    except PolicySchemaError as exc:
        _log_model_output_failure(flow=flow, raw_text=raw_text, error=exc, context=context)
        raise HTTPException(status_code=502, detail=f"AI returned invalid policy schema: {exc}") from exc


def _parse_generic_model_object(
    raw_text: str,
    *,
    flow: str,
    context: dict[str, object] | None = None,
) -> dict:
    try:
        parsed = parse_model_json(raw_text)
    except ParsedModelOutputError as exc:
        _log_model_output_failure(flow=flow, raw_text=raw_text, error=exc, context=context)
        raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {exc}") from exc
    if not isinstance(parsed, dict):
        error = PolicySchemaError("Model output must be a JSON object.")
        _log_model_output_failure(flow=flow, raw_text=raw_text, error=error, context=context)
        raise HTTPException(status_code=502, detail="AI returned an invalid JSON object.")
    return parsed


def _normalize_policy_payload_or_400(
    policy_data: dict,
    *,
    organization_hint: str,
    required_frameworks: list[str] | None = None,
) -> dict:
    try:
        return normalize_policy_payload(
            policy_data,
            organization_hint=organization_hint,
            required_frameworks=required_frameworks,
        )
    except PolicySchemaError as exc:
        raise HTTPException(status_code=400, detail=f"Policy data is invalid: {exc}") from exc


def _build_preview_text(policy_data: dict) -> str:
    sections = policy_data.get("sections")
    if isinstance(sections, list):
        for slot_id in ("purpose", "scope", "policy_statement"):
            for section in sections:
                if isinstance(section, dict) and str(section.get("slot_id")) == slot_id:
                    candidate = str(section.get("content") or "").strip()
                    if candidate:
                        return candidate.replace("\n", " ")[:280]
    for key in ("purpose", "scope", "policy_statement"):
        candidate = policy_data.get(key)
        if isinstance(candidate, str) and candidate.strip():
            return candidate.strip().replace("\n", " ")[:280]
    return "Document generated successfully."


def _safe_text(value) -> str:
    if value is None:
        return "-"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _write_temp_docx(docx_bytes: bytes) -> str:
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")
    with open(output_path, "wb") as f:
        f.write(docx_bytes)
    return output_path


def _file_docx_response(
    *,
    tenant_id: str,
    policy_data: dict,
    template_name: str,
    output_name: str,
    document_name: str,
    doc_type: str,
    source_name: str,
    policy_id: str | None = None,
    watermark_exports: bool = False,
) -> FileResponse:
    try:
        docx_bytes = _build_docx(policy_data, template_name)
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Document rendering error: {exc}") from exc

    if not docx_bytes:
        raise HTTPException(status_code=500, detail="Document rendering produced an empty file.")

    if watermark_exports:
        docx_bytes = _watermark_docx_bytes(docx_bytes)

    preview_text = _build_preview_text(policy_data)
    try:
        record = save_generated_document(
            workspace_id=tenant_id,
            filename=output_name,
            document_name=document_name,
            doc_type=doc_type,
            preview=preview_text,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_bytes=docx_bytes,
            source_name=source_name,
            policy_number=policy_data.get("policy_number"),
            version=policy_data.get("version"),
            policy_id=policy_id,
        )
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    temp_path = _write_temp_docx(docx_bytes)

    return FileResponse(
        path=temp_path,
        filename=output_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "X-Midnight-Preview": preview_text,
            "X-Midnight-Document-Id": record["id"],
        },
    )


class BirdsongRequest(BaseModel):
    messages: list[dict]
    system: Optional[str] = None


@pipeline_router.get("/smoke-docx")
async def smoke_docx():
    doc = Document()
    doc.add_paragraph("Hello world from Midnight Core")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    temp_path = _write_temp_docx(buffer.read())

    return FileResponse(
        path=temp_path,
        filename="smoke_test.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@pipeline_router.post("/birdsong")
async def birdsong(request: BirdsongRequest):
    try:
        client = _get_anthropic_client()
        system = request.system or (
            "You are Midnight, the AI compliance assistant for the Midnight compliance platform "
            "by Takeoff LLC. You are also known as Bird Talk, a chicken mascot who is the Chief "
            "Compliance Officer. Help users understand their compliance program, identify gaps, "
            "and guide them toward building the right policies and controls. "
            "Midnight covers HIPAA, HITRUST-aligned domains, PCI DSS, NIST CSF, and SOC 2. "
            "Philosophy: Handshake, not takeover - human-led, AI-accelerated. "
            "Keep responses short (3-5 sentences max). Warm, direct, expert tone."
        )

        message = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=8096,
            system=system,
            messages=request.messages,
        )
        return {"reply": message.content[0].text}
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Bird Talk error: {exc}") from exc


MIGRATION_SYSTEM_PROMPT = """You are a policy reconstruction engine for Midnight, Takeoff LLC's enterprise compliance platform.

You receive extracted content from a legacy policy document and must reconstruct it into a clean, structured policy.

Your output must be valid JSON with this exact structure:
{
  "policy_name": "...",
  "doc_type": "...",
  "version": "...",
  "effective_date": "...",
  "owner": "...",
  "purpose": "...",
  "scope": "...",
  "policy_statement": "...",
  "definitions": [...],
  "procedures": [...],
  "roles_responsibilities": [...],
  "exceptions": "...",
  "enforcement": "...",
  "references": [...],
  "revision_history": [...],
  "framework_mappings": {
    "<framework>": ["<control>", "..."]
  },
  "framework_map": {
    "overall_coverage": "<Strong|Moderate|Weak|Critical>",
    "total_controls_mapped": 0,
    "total_gaps": 0,
    "frameworks_covered": [],
    "audit_summary": "...",
    "mapped_citations": [],
    "gaps": [
      {
        "framework": "...",
        "control_id": "...",
        "gap_description": "...",
        "risk_level": "<high|medium|low>",
        "suggestion": "..."
      }
    ]
  },
  "gaps": [...],
  "quality_score": 0
}

Rules:
- Never invent content not in the source.
- For mapped source sections and header fields, copy content verbatim.
- Never paraphrase or summarize mapped source sections.
- Preserve bullets, indentation, and numbered lists exactly for mapped source sections.
- Replace semicolons in mapped body content with line breaks.
- Never append unmapped content to the end of the document.
- Do not restructure source tables.
- Map only to the provided framework controls or HITRUST-aligned domains.
- framework_map.gaps must list specific missing controls with risk level and remediation suggestion.
- quality_score is 0-100 based on completeness.
- Return only valid JSON, no markdown and no preamble."""


async def _extract_policy_data(
    *,
    file_bytes: bytes,
    filename: str,
    template_name: str,
    frameworks: list[str],
    template_section_map: Optional[dict] = None,
    organization_hint: str = "",
) -> dict:
    raw_text = _extract_text_from_upload(file_bytes, filename)
    if len(raw_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Document appears to be empty or unreadable.")
    mapped_source = _map_migrate_source(file_bytes, filename)
    mapped_context = _build_migrate_mapping_context(mapped_source)

    normalized_frameworks, fw_context = build_framework_prompt_context(frameworks)
    mapping_rules = build_framework_mapping_rules(frameworks)

    user_msg = (
        f"DOCUMENT: {filename}\n"
        f"TEMPLATE TYPE: {template_name}\n"
        f"TEMPLATE SECTION MAP:\n{json.dumps(template_section_map or _default_template_section_map(template_name), ensure_ascii=False)}\n\n"
        f"FRAMEWORKS TO MAP: {', '.join(normalized_frameworks)}\n"
        f"FRAMEWORK CONTROLS REFERENCE:\n{chr(10).join(fw_context)}\n\n"
        f"MAPPING RULES: {mapping_rules}\n\n"
        f"VERBATIM SOURCE MAPPING:\n---\n{mapped_context or 'No deterministic section matches found.'}\n---\n\n"
        f"SOURCE CONTENT:\n---\n{raw_text[:8000]}\n---\n\n"
        "Reconstruct this policy into the required JSON structure. Return only JSON."
    )

    client = _get_anthropic_client()
    message = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        system=MIGRATION_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}],
    )
    policy_data = _parse_policy_model_output(
        message.content[0].text,
        flow="migrate_policy",
        organization_hint=organization_hint,
        required_frameworks=normalized_frameworks,
        context={"filename": filename, "template_name": template_name},
    )
    return _apply_migrate_source_mapping(policy_data, mapped_source)


CREATE_SYSTEM_PROMPT = """You are a policy creation engine for Midnight, Takeoff LLC's enterprise compliance platform.

Create a complete, professional enterprise policy document based on the intake information provided.

Return valid JSON with this structure:
{
  "policy_name": "...",
  "doc_type": "...",
  "version": "1.0",
  "effective_date": "...",
  "owner": "...",
  "purpose": "...",
  "scope": "...",
  "policy_statement": "...",
  "definitions": [...],
  "procedures": [...],
  "roles_responsibilities": [...],
  "exceptions": "...",
  "enforcement": "...",
  "references": [...],
  "revision_history": [{"version": "1.0", "date": "...", "description": "Initial release"}],
  "framework_mappings": {
    "<framework>": ["<control>", "..."]
  },
  "framework_map": {
    "overall_coverage": "Strong",
    "total_controls_mapped": 0,
    "total_gaps": 0,
    "frameworks_covered": [],
    "audit_summary": "...",
    "mapped_citations": [],
    "gaps": []
  }
}

Rules:
- Write professionally and specifically for the industry provided.
- Map only to the provided framework controls or HITRUST-aligned domains.
- Procedures must be written step-by-step when procedural detail is appropriate.
- Return only valid JSON, no markdown."""


class CreatePolicyRequest(BaseModel):
    policy_name: str
    doc_type: str
    industry: str
    frameworks: list[str]
    owner: str
    description: Optional[str] = None
    policy_number: Optional[str] = None
    version: Optional[str] = None
    grc_id: Optional[str] = None
    effective_date: Optional[str] = None
    last_reviewed: Optional[str] = None
    last_revised: Optional[str] = None
    supersedes: Optional[str] = None
    custodians: Optional[str] = None
    owner_title: Optional[str] = None
    approver_name: Optional[str] = None
    approver_title: Optional[str] = None
    date_signed: Optional[str] = None
    date_approved: Optional[str] = None
    purpose_scope: Optional[str] = None
    definitions_text: Optional[str] = None
    policy_statement: Optional[str] = None
    procedures_text: Optional[str] = None
    related_policies: Optional[str] = None
    citations_references: Optional[str] = None


def _anthropic_text_response(message: Any) -> tuple[str, str]:
    chunks: list[str] = []
    for block in getattr(message, "content", []) or []:
        text = getattr(block, "text", None)
        if text:
            chunks.append(text)
    return "\n".join(chunks).strip(), str(getattr(message, "stop_reason", "") or "")


def _call_model_json_object(
    *,
    system_prompt: str,
    user_prompt: str,
    flow: str,
    context: dict[str, object] | None = None,
    max_tokens: int = 1200,
) -> dict[str, Any]:
    client = _get_anthropic_client()
    last_error: Exception | None = None
    last_raw_text = ""
    prompt = user_prompt

    for attempt in range(2):
        message = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": prompt}],
        )
        raw_text, stop_reason = _anthropic_text_response(message)
        last_raw_text = raw_text

        try:
            if stop_reason == "max_tokens":
                raise ParsedModelOutputError("AI response was truncated before completing valid JSON.")
            parsed = parse_model_json(raw_text)
            if not isinstance(parsed, dict):
                raise PolicySchemaError("Model output must be a JSON object.")
            return parsed
        except (ParsedModelOutputError, PolicySchemaError) as exc:
            last_error = exc
            if attempt == 0:
                prompt = (
                    f"{user_prompt}\n\n"
                    f"Your previous response was invalid because: {exc}\n"
                    "Return one compact JSON object only. Use double-quoted keys, escape internal quotes, and do not include markdown fences or prose."
                )
                continue
            _log_model_output_failure(flow=flow, raw_text=last_raw_text, error=exc, context=context)
            raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {exc}") from exc

    raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {last_error}")


def _validate_policy_metadata(
    metadata: dict[str, Any],
    *,
    request: CreatePolicyRequest,
    organization_hint: str,
    normalized_frameworks: list[str],
) -> dict[str, Any]:
    title = str(metadata.get("title") or request.policy_name).strip()
    organization = str(metadata.get("organization") or organization_hint).strip()
    owner = str(metadata.get("owner") or request.owner).strip()
    document_type = str(metadata.get("document_type") or request.doc_type or "Policy").strip()
    status = str(metadata.get("status") or "Draft").strip() or "Draft"
    schema_version = str(metadata.get("schema_version") or POLICY_SCHEMA_VERSION).strip() or POLICY_SCHEMA_VERSION
    selected_frameworks = metadata.get("selected_frameworks") or normalized_frameworks
    if not isinstance(selected_frameworks, list):
        raise HTTPException(status_code=502, detail="AI returned invalid metadata: selected_frameworks must be a list.")
    selected_frameworks = [str(item).strip() for item in selected_frameworks if str(item).strip()]

    required = {
        "title": title,
        "organization": organization,
        "owner": owner,
        "document_type": document_type,
        "status": status,
        "schema_version": schema_version,
    }
    missing = [key for key, value in required.items() if not value]
    if missing:
        raise HTTPException(status_code=502, detail=f"AI returned invalid metadata: missing {', '.join(missing)}.")

    return {
        "title": title,
        "organization": organization,
        "owner": owner,
        "document_type": document_type,
        "status": status,
        "selected_frameworks": selected_frameworks,
        "schema_version": schema_version,
        "policy_number": request.policy_number or "",
        "version": request.version or "1.0",
        "grc_id": request.grc_id or "",
        "effective_date": request.effective_date or "",
        "last_reviewed": request.last_reviewed or "",
        "last_revised": request.last_revised or "",
        "supersedes": request.supersedes or "",
        "custodians": request.custodians or "",
        "owner_title": request.owner_title or "",
        "approver_name": request.approver_name or "",
        "approver_title": request.approver_title or "",
        "date_signed": request.date_signed or "",
        "date_approved": request.date_approved or "",
        "description": request.description or "",
        "industry": request.industry,
    }


def _validate_generated_section(section: dict[str, Any], *, slot_spec: dict[str, str]) -> dict[str, Any]:
    slot_id = str(section.get("slot_id") or slot_spec["slot_id"]).strip()
    heading = str(section.get("heading") or section.get("title") or slot_spec["heading"]).strip()
    content = str(section.get("content") or section.get("body") or "").strip()

    if slot_id != slot_spec["slot_id"]:
        raise PolicySchemaError(f"Expected slot_id '{slot_spec['slot_id']}' but received '{slot_id}'.")
    if not heading:
        raise PolicySchemaError(f"Section '{slot_id}' is missing a heading.")
    if not content:
        raise PolicySchemaError(f"Section '{slot_id}' is missing content.")
    if len(content) > SECTION_CONTENT_LIMIT:
        raise PolicySchemaError(f"Section '{slot_id}' exceeded the maximum content length.")

    return {
        "slot_id": slot_id,
        "heading": heading,
        "content": content,
        "sort_order": POLICY_REQUIRED_SLOTS.index(slot_id) + 1,
        "source_origin": "ai_generated",
    }


def _build_policy_payload_from_sections(
    *,
    metadata: dict[str, Any],
    sections: list[dict[str, Any]],
    section_errors: list[dict[str, str]] | None = None,
) -> dict[str, Any]:
    section_map = {section["slot_id"]: section for section in sections}
    framework_mappings = {
        framework: [] for framework in metadata.get("selected_frameworks", [])
    } if metadata.get("selected_frameworks") else {}

    payload: dict[str, Any] = {
        "title": metadata["title"],
        "policy_name": metadata["title"],
        "organization": metadata["organization"],
        "owner": metadata["owner"],
        "doc_type": metadata["document_type"],
        "status": metadata["status"],
        "version": metadata.get("version") or "1.0",
        "policy_number": metadata.get("policy_number") or "",
        "grc_id": metadata.get("grc_id") or "",
        "effective_date": metadata.get("effective_date") or "",
        "last_reviewed": metadata.get("last_reviewed") or "",
        "last_revised": metadata.get("last_revised") or "",
        "supersedes": metadata.get("supersedes") or "",
        "custodians": metadata.get("custodians") or "",
        "owner_title": metadata.get("owner_title") or "",
        "approver_name": metadata.get("approver_name") or "",
        "approver_title": metadata.get("approver_title") or "",
        "date_signed": metadata.get("date_signed") or "",
        "date_approved": metadata.get("date_approved") or "",
        "metadata": {
            "owner": metadata["owner"],
            "document_type": metadata["document_type"],
            "schema_version": metadata["schema_version"],
            "selected_frameworks": metadata.get("selected_frameworks", []),
        },
        "sections": sections,
        "framework_mappings": framework_mappings,
        "framework_map": {
            "overall_coverage": "In progress" if framework_mappings else "Not selected",
            "total_controls_mapped": 0,
            "total_gaps": 0,
            "frameworks_covered": list(framework_mappings.keys()),
            "audit_summary": "Framework mappings will be refined as policy sections are reviewed.",
            "mapped_citations": [],
            "gaps": [],
        },
        "section_errors": section_errors or [],
    }

    for slot_spec in POLICY_SLOT_SPECS:
        payload[slot_spec["slot_id"]] = section_map.get(slot_spec["slot_id"], {}).get("content", "")
    return payload


def _ensure_required_slots(policy_data: dict[str, Any]) -> list[str]:
    sections = policy_data.get("sections") if isinstance(policy_data.get("sections"), list) else []
    present = {str(section.get("slot_id") or "").strip() for section in sections if isinstance(section, dict)}
    missing = [slot for slot in POLICY_REQUIRED_SLOTS if slot not in present]
    return missing


def _ensure_required_slots_or_400(policy_data: dict[str, Any]) -> dict[str, Any]:
    missing = _ensure_required_slots(policy_data)
    if missing:
        raise HTTPException(
            status_code=400,
            detail=f"Policy draft is incomplete. Missing required sections: {', '.join(missing)}.",
        )
    return policy_data


def _merge_sections_from_top_level(policy_data: dict[str, Any]) -> dict[str, Any]:
    section_map: dict[str, dict[str, Any]] = {}
    for section in policy_data.get("sections", []) if isinstance(policy_data.get("sections"), list) else []:
        if isinstance(section, dict):
            slot_id = str(section.get("slot_id") or "").strip()
            if slot_id:
                section_map[slot_id] = dict(section)

    for index, spec in enumerate(POLICY_SLOT_SPECS, start=1):
        top_level_value = policy_data.get(spec["slot_id"])
        if isinstance(top_level_value, str) and top_level_value.strip():
            section_map[spec["slot_id"]] = {
                "slot_id": spec["slot_id"],
                "heading": spec["heading"],
                "content": top_level_value.strip(),
                "sort_order": index,
            }
        elif spec["slot_id"] in section_map:
            section_map[spec["slot_id"]].setdefault("heading", spec["heading"])
            section_map[spec["slot_id"]].setdefault("sort_order", index)

    ordered_sections = [
        section_map[spec["slot_id"]]
        for spec in POLICY_SLOT_SPECS
        if spec["slot_id"] in section_map
    ]
    merged = dict(policy_data)
    merged["sections"] = ordered_sections
    return merged


def _build_metadata_prompt(
    request: CreatePolicyRequest,
    *,
    organization_hint: str,
    normalized_frameworks: list[str],
) -> str:
    return (
        f"Policy Name: {request.policy_name}\n"
        f"Requested Document Type: {request.doc_type}\n"
        f"Organization: {organization_hint or 'Not provided'}\n"
        f"Industry: {request.industry}\n"
        f"Owner: {request.owner}\n"
        f"Description: {request.description or 'Not provided'}\n"
        f"Frameworks: {', '.join(normalized_frameworks) or 'None selected'}\n\n"
        "Return valid JSON only with the fields: "
        "title, organization, owner, document_type, status, selected_frameworks, schema_version."
    )


def _build_section_prompt(
    request: CreatePolicyRequest,
    *,
    metadata: dict[str, Any],
    slot_spec: dict[str, str],
    normalized_frameworks: list[str],
    fw_context: list[str],
    mapping_rules: str,
) -> str:
    return (
        f"Policy Title: {metadata['title']}\n"
        f"Organization: {metadata['organization']}\n"
        f"Owner: {metadata['owner']}\n"
        f"Industry: {request.industry}\n"
        f"Frameworks: {', '.join(normalized_frameworks) or 'None selected'}\n"
        f"Document Type: {metadata['document_type']}\n"
        f"Policy Description: {request.description or 'Not provided'}\n"
        f"Requested Slot: {slot_spec['slot_id']}\n"
        f"Required Heading: {slot_spec['heading']}\n"
        f"Section Guidance: {slot_spec['instruction']}\n"
        f"Framework Controls Reference:\n{chr(10).join(fw_context)}\n\n"
        f"Mapping Rules: {mapping_rules}\n\n"
        "Generate only one policy section. Return valid JSON only in this shape:\n"
        '{\n  "slot_id": "' + slot_spec["slot_id"] + '",\n  "heading": "' + slot_spec["heading"] + '",\n  "content": "..."\n}\n'
        "Do not include markdown fences or prose."
    )


async def _generate_policy_data(
    request: CreatePolicyRequest,
    *,
    tenant_id: str,
    organization_hint: str = "",
    existing_policy_id: str | None = None,
) -> tuple[dict[str, Any], str]:
    normalized_frameworks, fw_context = build_framework_prompt_context(request.frameworks)
    mapping_rules = build_framework_mapping_rules(request.frameworks)

    metadata_raw = _call_model_json_object(
        system_prompt="You are Midnight's policy metadata generator. Return JSON only.",
        user_prompt=_build_metadata_prompt(request, organization_hint=organization_hint, normalized_frameworks=normalized_frameworks),
        flow="create_policy_metadata",
        context={"policy_name": request.policy_name, "doc_type": request.doc_type},
        max_tokens=700,
    )
    metadata = _validate_policy_metadata(
        metadata_raw,
        request=request,
        organization_hint=organization_hint,
        normalized_frameworks=normalized_frameworks,
    )

    draft_record = save_policy_draft(
        tenant_id=tenant_id,
        title=metadata["title"],
        document_type=metadata["document_type"],
        organization=metadata["organization"],
        owner=metadata["owner"],
        status=metadata["status"],
        schema_version=metadata["schema_version"],
        selected_frameworks=metadata["selected_frameworks"],
        sections=[],
        policy_number=metadata.get("policy_number") or None,
        version=metadata.get("version") or None,
        policy_id=existing_policy_id,
    )
    policy_id = draft_record["policy"]["id"]

    sections: list[dict[str, Any]] = []
    section_errors: list[dict[str, str]] = []

    for slot_spec in POLICY_SLOT_SPECS:
        try:
            raw_section = _call_model_json_object(
                system_prompt="You are Midnight's policy section generator. Return JSON only for one section.",
                user_prompt=_build_section_prompt(
                    request,
                    metadata=metadata,
                    slot_spec=slot_spec,
                    normalized_frameworks=normalized_frameworks,
                    fw_context=fw_context,
                    mapping_rules=mapping_rules,
                ),
                flow=f"create_policy_section_{slot_spec['slot_id']}",
                context={"policy_name": request.policy_name, "slot_id": slot_spec["slot_id"]},
                max_tokens=1100,
            )
            section = _validate_generated_section(raw_section, slot_spec=slot_spec)
            sections.append(section)
            save_policy_draft(
                tenant_id=tenant_id,
                title=metadata["title"],
                document_type=metadata["document_type"],
                organization=metadata["organization"],
                owner=metadata["owner"],
                status=metadata["status"],
                schema_version=metadata["schema_version"],
                selected_frameworks=metadata["selected_frameworks"],
                sections=sections,
                policy_number=metadata.get("policy_number") or None,
                version=metadata.get("version") or None,
                policy_id=policy_id,
            )
        except HTTPException as exc:
            section_errors.append(
                {
                    "slot_id": slot_spec["slot_id"],
                    "heading": slot_spec["heading"],
                    "error": str(exc.detail),
                }
            )

    policy_payload = _build_policy_payload_from_sections(
        metadata=metadata,
        sections=sections,
        section_errors=section_errors,
    )
    policy_payload["_draft"] = {"policy_id": policy_id}
    return policy_payload, policy_id


def _build_docx(policy_data: dict, template_name: str) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        return heading

    def add_body(text: str):
        paragraph = doc.add_paragraph(_safe_text(text))
        return paragraph

    def add_bullet(text: str):
        doc.add_paragraph(_safe_text(text), style="List Bullet")

    fields = [
        ("Policy Name", policy_data.get("policy_name", "Untitled Policy")),
        ("Document Type", policy_data.get("doc_type", "Policy")),
        ("Version", policy_data.get("version", "1.0")),
        ("Effective Date", policy_data.get("effective_date", "-")),
        ("Policy Owner", policy_data.get("owner", "-")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"

    for index, (label, value) in enumerate(fields):
        row = table.rows[index]
        row.cells[0].text = _safe_text(label)
        row.cells[1].text = _safe_text(value)

    doc.add_paragraph()

    rendered_sections = policy_data.get("sections")
    if isinstance(rendered_sections, list) and rendered_sections:
        sections = [
            (f"{index}. {section.get('heading') or 'Section'}", str(section.get("slot_id") or f"section_{index}"))
            for index, section in enumerate(rendered_sections, start=1)
        ]
        section_content_map = {
            str(section.get("slot_id") or f"section_{index}"): section.get("content")
            for index, section in enumerate(rendered_sections, start=1)
            if isinstance(section, dict)
        }
    else:
        sections = [
            ("1. Purpose", "purpose"),
            ("2. Scope", "scope"),
            ("3. Policy Statement", "policy_statement"),
            ("4. Definitions", "definitions"),
            ("5. Procedures", "procedures"),
            ("6. Roles & Responsibilities", "roles_responsibilities"),
            ("7. Exceptions", "exceptions"),
            ("8. Enforcement", "enforcement"),
            ("9. References", "references"),
        ]
        section_content_map = {key: policy_data.get(key) for _, key in sections}

    for title, key in sections:
        add_heading(title, level=1)
        content = section_content_map.get(key)

        if not content:
            add_body("[Section not found in source document]")
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    label = item.get("role") or item.get("title") or ""
                    desc = item.get("responsibility") or item.get("description") or item
                    text = f"{label}: {desc}" if label else _safe_text(desc)
                    add_bullet(text)
                else:
                    add_bullet(item)
        else:
            add_body(content)

        doc.add_paragraph()

    framework_mappings = policy_data.get("framework_mappings", {})
    if framework_mappings:
        add_heading("10. Framework Mappings", level=1)
        for framework, controls in framework_mappings.items():
            if not controls:
                continue
            add_heading(_safe_text(framework), level=2)
            if isinstance(controls, list):
                for control in controls:
                    add_bullet(control)
            else:
                add_body(controls)
        doc.add_paragraph()

    framework_map = policy_data.get("framework_map", {}) or {}
    gaps = framework_map.get("gaps") or policy_data.get("gaps", [])
    if gaps:
        add_heading("11. Gap Analysis", level=1)
        for gap in gaps:
            if isinstance(gap, dict):
                label = (
                    f"{gap.get('framework', '')} {gap.get('control_id', '')}: "
                    f"{gap.get('gap_description', '')}"
                ).strip()
                if gap.get("suggestion"):
                    label = f"{label} -> {gap['suggestion']}"
                add_bullet(label)
            else:
                add_bullet(gap)
        doc.add_paragraph()

    revision_history = policy_data.get("revision_history", [])
    add_heading("12. Revision History", level=1)
    if revision_history:
        rev_table = doc.add_table(rows=1, cols=3)
        rev_table.style = "Table Grid"
        header = rev_table.rows[0].cells
        header[0].text = "Version"
        header[1].text = "Date"
        header[2].text = "Description"

        for entry in revision_history:
            row = rev_table.add_row().cells
            if isinstance(entry, dict):
                row[0].text = _safe_text(entry.get("version", ""))
                row[1].text = _safe_text(entry.get("date", ""))
                row[2].text = _safe_text(entry.get("description", entry.get("change", "")))
            else:
                row[0].text = _safe_text(entry)
                row[1].text = ""
                row[2].text = ""
    else:
        add_body("[No revision history found in source document]")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = (
        f"{policy_data.get('policy_name', 'Policy')} · "
        f"v{policy_data.get('version', '1.0')} · "
        "Midnight - Takeoff LLC · CONFIDENTIAL"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@pipeline_router.post("/migrate/preview")
async def migrate_preview(
    request: Request,
    source_file: UploadFile = File(...),
    template_file: UploadFile | None = File(default=None),
    template_name: str = Form(default="generic_policy"),
    industry: str = Form(default="Healthcare"),
    frameworks: str = Form(default="HIPAA,HITRUST"),
):
    _require_supported_file(source_file, "source_file")
    file_bytes = await source_file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    framework_list = [item.strip() for item in frameworks.split(",") if item.strip()]
    tenant = _enforce_trial_limits(request, frameworks=framework_list, upload_action="migrate_upload")
    template_section_map = _default_template_section_map(template_name)
    template_reference_name = template_name

    if template_file is not None:
        _require_supported_template_file(template_file, "template_file")
        template_file_bytes = await template_file.read()
        if not template_file_bytes:
            raise HTTPException(status_code=400, detail="Template reference file is empty.")
        template_reference_name = template_file.filename or template_name
        template_section_map = _extract_template_section_map(
            template_name=template_name,
            template_file_bytes=template_file_bytes,
            template_filename=template_reference_name,
        )

    try:
        policy_data = await _extract_policy_data(
            file_bytes=file_bytes,
            filename=source_file.filename or "document.txt",
            template_name=template_name,
            frameworks=framework_list,
            template_section_map=template_section_map,
            organization_hint=str(tenant.get("name") or ""),
        )
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Extraction error: {exc}") from exc

    create_activity_event(tenant_id=tenant["id"], action="migrate_upload")

    policy_data["_session"] = {
        "source_filename": source_file.filename or "document",
        "template_name": template_name,
        "template_reference_name": template_reference_name,
        "frameworks": framework_list,
        "industry": industry,
        "preview_id": str(uuid.uuid4()),
        "tenant_id": tenant["id"],
        "plan_type": tenant.get("plan_type", "trial"),
    }
    policy_data["template_section_map"] = template_section_map

    return JSONResponse(
        content={
            "policy_data": policy_data,
            "framework_map": policy_data.get("framework_map"),
            "quality_score": policy_data.get("quality_score", 0),
            "template_section_map": template_section_map,
            "preview_id": policy_data["_session"]["preview_id"],
        }
    )


class MigrateGenerateRequest(BaseModel):
    policy_data: dict


@pipeline_router.post("/migrate/generate")
async def migrate_generate(request: Request, payload: MigrateGenerateRequest):
    policy_data = dict(payload.policy_data or {})
    if not policy_data:
        raise HTTPException(status_code=400, detail="policy_data is required.")

    session = policy_data.pop("_session", {})
    source_name = session.get("source_filename", "document")
    template_name = session.get("template_name", "generic_policy")
    base_name = source_name.rsplit(".", 1)[0]
    output_name = f"{base_name}_migrated.docx"
    tenant = _tenant_context_from_request(request)
    policy_data = _normalize_policy_payload_or_400(
        policy_data,
        organization_hint=str(tenant.get("name") or ""),
        required_frameworks=[item for item in session.get("frameworks", []) if str(item).strip()],
    )

    return _file_docx_response(
        tenant_id=tenant["id"],
        policy_data=policy_data,
        template_name=template_name,
        output_name=output_name,
        document_name=policy_data.get("policy_name", base_name),
        doc_type=policy_data.get("doc_type", template_name),
        source_name=source_name,
        watermark_exports=str(tenant.get("plan_type") or "").lower() == "trial",
    )


@pipeline_router.post("/migrate")
async def migrate_document(
    request: Request,
    file: UploadFile | None = File(default=None),
    source_file: UploadFile | None = File(default=None),
    template_file: UploadFile | None = File(default=None),
    template_name: str = Form(default="generic_policy"),
    industry: str = Form(default="Healthcare"),
    frameworks: str = Form(default="HIPAA,HITRUST"),
):
    upload = file or source_file
    if upload is None:
        raise HTTPException(status_code=400, detail="No document was uploaded.")

    _require_supported_file(upload)
    file_bytes = await upload.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    framework_list = [item.strip() for item in frameworks.split(",") if item.strip()]
    tenant = _enforce_trial_limits(request, frameworks=framework_list, upload_action="migrate_upload")
    template_section_map = _default_template_section_map(template_name)

    if template_file is not None:
        _require_supported_template_file(template_file, "template_file")
        template_file_bytes = await template_file.read()
        if not template_file_bytes:
            raise HTTPException(status_code=400, detail="Template reference file is empty.")
        template_section_map = _extract_template_section_map(
            template_name=template_name,
            template_file_bytes=template_file_bytes,
            template_filename=template_file.filename or template_name,
        )

    try:
        policy_data = await _extract_policy_data(
            file_bytes=file_bytes,
            filename=upload.filename or "document.txt",
            template_name=template_name,
            frameworks=framework_list,
            template_section_map=template_section_map,
            organization_hint=str(tenant.get("name") or ""),
        )
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Migration error: {exc}") from exc

    create_activity_event(tenant_id=tenant["id"], action="migrate_upload")

    base_name = (upload.filename or "document").rsplit(".", 1)[0]
    return _file_docx_response(
        tenant_id=tenant["id"],
        policy_data=policy_data,
        template_name=template_name,
        output_name=f"{base_name}_migrated.docx",
        document_name=policy_data.get("policy_name", base_name),
        doc_type=policy_data.get("doc_type", template_name),
        source_name=upload.filename or "document",
        watermark_exports=str(tenant.get("plan_type") or "").lower() == "trial",
    )


@pipeline_router.post("/create/preview")
async def create_preview(request: Request, payload: CreatePolicyRequest):
    try:
        framework_list = [item.strip() for item in payload.frameworks if item.strip()]
        tenant = _enforce_trial_limits(request, frameworks=framework_list)
        policy_data, policy_id = await _generate_policy_data(
            payload,
            tenant_id=tenant["id"],
            organization_hint=str(tenant.get("name") or ""),
        )
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Creation error: {exc}") from exc

    policy_data["_session"] = {
        "doc_type": payload.doc_type,
        "industry": payload.industry,
        "frameworks": payload.frameworks,
        "source_name": payload.policy_name,
        "preview_id": str(uuid.uuid4()),
        "tenant_id": tenant["id"],
        "plan_type": tenant.get("plan_type", "trial"),
        "policy_id": policy_id,
    }

    return JSONResponse(
        content={
            "policy_data": policy_data,
            "framework_map": policy_data.get("framework_map"),
            "preview_id": policy_data["_session"]["preview_id"],
            "section_errors": policy_data.get("section_errors", []),
        }
    )


class CreateGenerateRequest(BaseModel):
    policy_data: dict


@pipeline_router.post("/create/generate")
async def create_generate(request: Request, payload: CreateGenerateRequest):
    policy_data = dict(payload.policy_data or {})
    if not policy_data:
        raise HTTPException(status_code=400, detail="policy_data is required.")

    session = policy_data.pop("_session", {})
    policy_name = policy_data.get("policy_name", session.get("source_name", "Policy"))
    doc_type = policy_data.get("doc_type", session.get("doc_type", "POLICY"))
    output_name = f"{policy_name.replace(' ', '_')}_v{policy_data.get('version', '1.0')}.docx"
    tenant = _tenant_context_from_request(request)
    policy_data = _merge_sections_from_top_level(policy_data)
    policy_data = _normalize_policy_payload_or_400(
        policy_data,
        organization_hint=str(tenant.get("name") or ""),
        required_frameworks=[item for item in session.get("frameworks", []) if str(item).strip()],
    )
    policy_data = _ensure_required_slots_or_400(policy_data)
    try:
        save_policy_draft(
            tenant_id=tenant["id"],
            title=policy_data.get("title") or policy_name,
            document_type=policy_data.get("doc_type") or doc_type,
            organization=policy_data.get("organization") or str(tenant.get("name") or ""),
            owner=policy_data.get("owner") or "Unknown",
            status=policy_data.get("status") or "Draft",
            schema_version=(policy_data.get("metadata") or {}).get("schema_version", POLICY_SCHEMA_VERSION),
            selected_frameworks=[item for item in session.get("frameworks", []) if str(item).strip()],
            sections=policy_data.get("sections", []),
            policy_number=policy_data.get("policy_number") or None,
            version=policy_data.get("version") or None,
            policy_id=session.get("policy_id"),
        )
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return _file_docx_response(
        tenant_id=tenant["id"],
        policy_data=policy_data,
        template_name=str(doc_type).lower(),
        output_name=output_name,
        document_name=policy_name,
        doc_type=doc_type,
        source_name=session.get("source_name", policy_name),
        policy_id=session.get("policy_id"),
        watermark_exports=str(tenant.get("plan_type") or "").lower() == "trial",
    )


@pipeline_router.post("/create")
async def create_document(request: Request, payload: CreatePolicyRequest):
    try:
        framework_list = [item.strip() for item in payload.frameworks if item.strip()]
        tenant = _enforce_trial_limits(request, frameworks=framework_list)
        policy_data, policy_id = await _generate_policy_data(
            payload,
            tenant_id=tenant["id"],
            organization_hint=str(tenant.get("name") or ""),
        )
        policy_data = _ensure_required_slots_or_400(policy_data)
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Creation error: {exc}") from exc

    output_name = f"{payload.policy_name.replace(' ', '_')}_v1.0.docx"
    return _file_docx_response(
        tenant_id=tenant["id"],
        policy_data=policy_data,
        template_name=payload.doc_type,
        output_name=output_name,
        document_name=policy_data.get("policy_name", payload.policy_name),
        doc_type=payload.doc_type,
        source_name=payload.policy_name,
        policy_id=policy_id,
        watermark_exports=str(tenant.get("plan_type") or "").lower() == "trial",
    )


class GrcSummaryRequest(BaseModel):
    organization_name: str
    industry: str
    frameworks: list[str]
    policy_data: Optional[dict] = None


class TenantCreateRequest(BaseModel):
    company_name: str
    industry: Optional[str] = None
    region: Optional[str] = None
    employee_count: Optional[str] = None


class OnboardingUpdateRequest(BaseModel):
    current_step: Optional[str] = None
    progress: Optional[int] = None
    build_method: Optional[str] = None
    primary_objective: Optional[str] = None
    frameworks: Optional[list[str]] = None
    enabled_modules: Optional[list[str]] = None
    completed: Optional[bool] = None


class TenantActivateRequest(BaseModel):
    enabled_modules: list[str] = []


@pipeline_router.post("/grc-summary")
async def create_grc_summary(request: Request, payload: GrcSummaryRequest):
    normalized_frameworks, _ = build_framework_prompt_context(payload.frameworks)
    tenant = _enforce_trial_limits(request, frameworks=normalized_frameworks)
    try:
        documents = list_generated_documents(tenant["id"])
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    pdf_bytes = build_grc_summary_pdf(
        organization_name=payload.organization_name.strip() or "Organization",
        industry=payload.industry.strip() or "Unspecified",
        frameworks=normalized_frameworks,
        documents=documents,
    )

    if str(tenant.get("plan_type") or "").lower() == "trial":
        pdf_bytes = build_grc_summary_pdf(
            organization_name=f"{TRIAL_WATERMARK_TEXT} - {payload.organization_name.strip() or 'Organization'}",
            industry=payload.industry.strip() or "Unspecified",
            frameworks=normalized_frameworks,
            documents=documents,
        )

    org_slug = (payload.organization_name.strip() or "organization").replace(" ", "_")
    output_name = f"{org_slug}_grc_summary.pdf"
    preview_text = (
        f"GRC summary for {payload.organization_name.strip() or 'your workspace'} "
        f"covering {', '.join(normalized_frameworks) or 'selected frameworks'}."
    )
    try:
        record = save_generated_document(
            workspace_id=tenant["id"],
            filename=output_name,
            document_name=f"{payload.organization_name.strip() or 'Organization'} GRC Summary",
            doc_type="PDF",
            preview=preview_text,
            content_type="application/pdf",
            file_bytes=pdf_bytes,
            source_name=payload.organization_name.strip() or output_name,
        )
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{output_name}"',
            "X-Midnight-Preview": preview_text,
            "X-Midnight-Document-Id": record["id"],
        },
    )


@pipeline_router.post("/analyze")
async def analyze_document(
    request: Request,
    file: UploadFile = File(...),
    frameworks: str = Form(default="HIPAA,PCI DSS,NIST CSF"),
):
    _require_supported_file(file)

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    framework_list = [item.strip() for item in frameworks.split(",") if item.strip()]
    _enforce_trial_limits(request, frameworks=framework_list, upload_action="analysis_upload")
    raw_text = _extract_text_from_upload(file_bytes, file.filename or "document")[:10000]

    normalized_frameworks, fw_context = build_framework_prompt_context(framework_list)
    mapping_rules = build_framework_mapping_rules(framework_list)

    user_msg = (
        f"FRAMEWORKS: {', '.join(normalized_frameworks)}\n\n"
        f"FRAMEWORK CONTROLS:\n{chr(10).join(fw_context)}\n\n"
        f"MAPPING RULES: {mapping_rules}\n\n"
        f"POLICY DOCUMENT:\n---\n{raw_text}\n---\n\n"
        "Analyze this policy against the framework controls. "
        "Return JSON with: covered_controls, missing_controls, partial_controls, gaps, recommendations, overall_score (0-100)"
    )

    try:
        client = _get_anthropic_client()
        message = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=2048,
            system="You are a GRC analyst. Analyze policy documents against compliance frameworks. Return only valid JSON.",
            messages=[{"role": "user", "content": user_msg}],
        )
        result = _parse_generic_model_object(
            message.content[0].text,
            flow="framework_analysis",
            context={"filename": file.filename or "document"},
        )
        create_activity_event(tenant_id=_tenant_id_from_request(request), action="analysis_upload")
        go_coverage = _go_framework_coverage(raw_text, normalized_frameworks, file.filename or "")
        if go_coverage:
            result["framework_coverage"] = go_coverage
        return JSONResponse(content=result)
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Analysis error: {exc}") from exc


@org_router.post("")
async def create_org(request: Request, payload: TenantCreateRequest):
    current_tenant_id = getattr(request.state, "tenant_id", None)
    if current_tenant_id:
        tenant = _tenant_context_from_request(request)
        return {
            "tenant_id": tenant["id"],
            "org_id": tenant["id"],
            "slug": tenant.get("slug"),
            "created": False,
        }

    slug = _slugify_value(payload.company_name)
    try:
        existing = get_tenant_by_slug(slug)
        if existing is not None:
            raise HTTPException(status_code=409, detail="A tenant with this slug already exists.")

        tenant = create_tenant(
            name=payload.company_name.strip(),
            slug=slug,
            industry=payload.industry,
            region=payload.region,
            employee_count=payload.employee_count,
            plan_type="trial",
        )
        update_profile_membership(
            user_id=getattr(request.state, "user_id", ""),
            tenant_id=tenant["id"],
            organization_name=payload.company_name.strip(),
        )
        create_onboarding_session(tenant_id=tenant["id"])
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return {
        "tenant_id": tenant["id"],
        "org_id": tenant["id"],
        "slug": tenant.get("slug"),
        "created": True,
    }


@org_router.patch("/{tenant_id}/onboarding")
async def update_org_onboarding(request: Request, tenant_id: str, payload: OnboardingUpdateRequest):
    tenant = _assert_tenant_access(request, tenant_id)
    frameworks = [item.strip() for item in (payload.frameworks or []) if item and item.strip()]
    if payload.frameworks is not None:
        _enforce_trial_limits(request, frameworks=frameworks)

    updates: dict[str, object] = {}
    for field in ("current_step", "progress", "build_method", "primary_objective", "completed"):
        value = getattr(payload, field)
        if value is not None:
            updates[field] = value
    if payload.frameworks is not None:
        updates["frameworks"] = frameworks
    if payload.enabled_modules is not None:
        updates["enabled_modules"] = [item.strip() for item in payload.enabled_modules if item and item.strip()]

    try:
        session = update_onboarding_session(tenant["id"], updates)
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return session


@org_router.post("/{tenant_id}/activate")
async def activate_org(request: Request, tenant_id: str, payload: TenantActivateRequest):
    tenant = _assert_tenant_access(request, tenant_id)
    modules = [item.strip() for item in payload.enabled_modules if item and item.strip()]

    try:
        session = update_onboarding_session(
            tenant["id"],
            {
                "completed": True,
                "current_step": "complete",
                "progress": 100,
                "enabled_modules": modules,
            },
        )
        enabled_modules = replace_enabled_modules(tenant["id"], modules)
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return {
        "tenant_id": tenant["id"],
        "org_id": tenant["id"],
        "slug": tenant.get("slug"),
        "enabled_modules": enabled_modules,
        "dashboard_config_url": f"/api/orgs/{tenant.get('slug')}/dashboard-config",
        "redirect_to": f"/{tenant.get('slug')}/dashboard",
        "onboarding_session": session,
    }


@org_router.get("/{org_slug}/dashboard-config")
async def get_dashboard_config(request: Request, org_slug: str):
    tenant = _tenant_context_from_request(request)
    if tenant.get("slug") != org_slug:
        raise HTTPException(status_code=403, detail="Cross-tenant access is not allowed.")

    try:
        session = get_onboarding_session(tenant["id"])
    except SupabaseStoreError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    return {
        "tenant_id": tenant["id"],
        "slug": tenant.get("slug"),
        "organization_name": tenant.get("name"),
        "plan_type": tenant.get("plan_type"),
        "industry": tenant.get("industry"),
        "region": tenant.get("region"),
        "employee_count": tenant.get("employee_count"),
        "frameworks": (session or {}).get("frameworks", []),
        "enabled_modules": (session or {}).get("enabled_modules", []),
        "build_method": (session or {}).get("build_method"),
        "primary_objective": (session or {}).get("primary_objective"),
        "completed": bool((session or {}).get("completed")),
    }


router.include_router(pipeline_router)
router.include_router(org_router)
