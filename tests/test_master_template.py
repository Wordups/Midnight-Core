"""The designed reusable master template: front matter, live TOC, page numbers,
tenant accent, and clean local-editing structure (real named styles)."""
import io
import os
import unittest
import zipfile

os.environ.setdefault("ANTHROPIC_API_KEY", "sk-test")
os.environ.setdefault("SUPABASE_URL", "https://example.supabase.co")
os.environ.setdefault("SUPABASE_ANON_KEY", "test")
os.environ.setdefault("SUPABASE_SERVICE_ROLE_KEY", "test")
os.environ.setdefault("ENVIRONMENT", "dev")

from docx import Document
from docx.shared import RGBColor

from backend.core.master_template import build_master_shell
from backend.core.template_engine import fill_template_placeholders


class MasterShellTests(unittest.TestCase):
    def test_shell_flags_and_structure(self):
        doc = build_master_shell("POLICY", "formal")
        self.assertTrue(doc.is_template_shell)
        self.assertTrue(doc.is_master_template)
        self.assertEqual(doc.section_heading_level, 1)
        lines = [p.text.strip() for p in doc.paragraphs]
        self.assertIn("POLICY", lines)                 # category eyebrow
        self.assertIn("Contents", lines)               # TOC label
        self.assertIn("Document Control", lines)

    def test_live_toc_field_present(self):
        doc = build_master_shell("POLICY", "formal")
        xml = doc.element.body.xml
        self.assertIn("TOC", xml)
        self.assertIn("instrText", xml)

    def test_named_styles_exist_for_local_editing(self):
        # Edits the user makes in Word must inherit the design via named styles.
        doc = build_master_shell("STANDARD", "modern")
        names = {s.name for s in doc.styles}
        for required in ("Heading 1", "Heading 2", "Heading 3", "Title", "Normal"):
            self.assertIn(required, names)

    def test_accent_applied_to_heading_one(self):
        doc = build_master_shell("POLICY", "formal", accent=RGBColor(0x22, 0x33, 0x44))
        self.assertEqual(doc.styles["Heading 1"].font.color.rgb, RGBColor(0x22, 0x33, 0x44))

    def test_category_label_varies(self):
        pol = build_master_shell("POLICY")
        run = build_master_shell("INCIDENT_RUNBOOK")
        self.assertIn("POLICY", [p.text.strip() for p in pol.paragraphs])
        self.assertIn("INCIDENT RUNBOOK", [p.text.strip() for p in run.paragraphs])


class MasterViaBuildDocxTests(unittest.TestCase):
    def _build(self):
        from backend.api.routes import _build_docx
        pd = {
            "doc_type": "POLICY", "template_variant": "formal",
            "policy_name": "Remote Access Security Policy",
            "organization": "Pixel Forge Interactive", "owner": "Riley Chen",
            "version": "2.0", "effective_date": "2026-08-13",
            "_branding": {"classification": "Confidential", "primary_color": "#7C5CD1"},
            "sections": [
                {"slot_id": "purpose", "heading": "Purpose", "content": "Body.\n\n- a\n- b"},
                {"slot_id": "scope", "heading": "Scope", "content": "Applies broadly."},
            ],
        }
        return _build_docx(pd, "generic")

    def test_no_raw_tokens_anywhere(self):
        raw = self._build()
        z = zipfile.ZipFile(io.BytesIO(raw))
        for name in z.namelist():
            if name.endswith(".xml"):
                self.assertNotIn("{{", z.read(name).decode("utf-8", "ignore"),
                                 msg=f"raw token leaked in {name}")

    def test_footer_has_page_numbers(self):
        raw = self._build()
        z = zipfile.ZipFile(io.BytesIO(raw))
        footer = z.read("word/footer1.xml").decode("utf-8", "ignore")
        self.assertIn("PAGE", footer)
        self.assertIn("NUMPAGES", footer)

    def test_sections_are_heading_one_and_numbered(self):
        raw = self._build()
        d = Document(io.BytesIO(raw))
        h1 = [p.text.strip() for p in d.paragraphs
              if p.style is not None and p.style.name == "Heading 1"]
        self.assertTrue(any(t.startswith("1. Purpose") for t in h1))
        self.assertTrue(any(t.startswith("2. Scope") for t in h1))
        # numbering unique (no duplicate "10." bug)
        import re
        nums = [t.split(".")[0] for t in h1 if re.match(r"^\d+\.", t)]
        self.assertEqual(len(nums), len(set(nums)))


if __name__ == "__main__":
    unittest.main()
