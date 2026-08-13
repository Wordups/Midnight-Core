"""Template engine: styled shells resolve per doc_type x variant, bodies are
cleared, branding lands, and everything fails open to a blank Document."""
import os
import unittest

os.environ.setdefault("ANTHROPIC_API_KEY", "sk-test")
os.environ.setdefault("SUPABASE_URL", "https://example.supabase.co")
os.environ.setdefault("SUPABASE_ANON_KEY", "test-anon")
os.environ.setdefault("SUPABASE_SERVICE_ROLE_KEY", "test-service-role")
os.environ.setdefault("ENVIRONMENT", "dev")

from docx.shared import RGBColor

from backend.core.template_engine import (
    DOC_TYPE_TO_CATEGORY,
    VARIANTS,
    load_template_shell,
    normalize_variant,
    shell_path,
)


class ShellResolutionTests(unittest.TestCase):
    def test_all_live_doc_types_have_all_variants(self):
        for doc_type in ("POLICY", "SOP", "STANDARD", "INCIDENT_RUNBOOK"):
            for variant in VARIANTS:
                path = shell_path(doc_type, variant)
                self.assertIsNotNone(path, f"{doc_type}/{variant} shell missing")

    def test_every_registered_category_resolves(self):
        for doc_type in DOC_TYPE_TO_CATEGORY:
            self.assertIsNotNone(shell_path(doc_type, "modern"), doc_type)

    def test_unknown_doc_type_returns_none(self):
        self.assertIsNone(shell_path("KARAOKE_NIGHT"))

    def test_variant_normalization(self):
        self.assertEqual(normalize_variant("FORMAL "), "formal")
        self.assertEqual(normalize_variant("nonsense"), "modern")
        self.assertEqual(normalize_variant(None), "modern")


class ShellLoadTests(unittest.TestCase):
    def test_shell_keeps_front_matter_and_styles(self):
        # Front matter (cover page, Document Control, TOC) is preserved; only the
        # sample sections are dropped. Styles survive.
        doc = load_template_shell("POLICY", "formal")
        self.assertTrue(doc.is_template_shell)
        self.assertGreater(len(doc.paragraphs), 0)
        style_names = {s.name for s in doc.styles}
        self.assertIn("Heading 1", style_names)
        self.assertIn("Title", style_names)

    def test_unknown_type_falls_back_to_blank(self):
        doc = load_template_shell("KARAOKE_NIGHT")
        self.assertFalse(doc.is_template_shell)

    def test_branding_color_applied_to_headings(self):
        doc = load_template_shell("POLICY", "modern", branding={
            "organization": "Pixel Forge Interactive",
            "title": "Access Control Policy",
            "primary_color": "#2563EB",
        })
        self.assertEqual(doc.styles["Heading 1"].font.color.rgb, RGBColor(0x25, 0x63, 0xEB))

    def test_bad_color_ignored(self):
        doc = load_template_shell("POLICY", "modern", branding={"primary_color": "not-a-color"})
        self.assertTrue(doc.is_template_shell)

    def test_render_into_shell_produces_valid_docx(self):
        from io import BytesIO
        from backend.renderers.docx_renderer import render_markdown_into

        doc = load_template_shell("SOP", "detailed")
        doc.add_heading("Purpose", level=1)
        render_markdown_into(doc, "This SOP defines **the steps** for onboarding.")
        buffer = BytesIO()
        doc.save(buffer)
        self.assertGreater(len(buffer.getvalue()), 1000)


if __name__ == "__main__":
    unittest.main()


class TableGridTests(unittest.TestCase):
    def test_shell_supports_table_grid_style(self):
        """The metadata/revision tables style themselves as Table Grid;
        pandoc shells lack it — the engine must register it (prod 500)."""
        doc = load_template_shell("POLICY", "formal")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"  # must not raise
        self.assertIsNotNone(table.style)


class PreviewAndLogoTests(unittest.TestCase):
    def test_preview_path_resolves_for_known_pairs(self):
        from backend.core.template_engine import preview_path
        self.assertIsNotNone(preview_path("policy", "formal"))
        self.assertIsNotNone(preview_path("audit_package", "executive"))

    def test_preview_path_rejects_traversal(self):
        from backend.core.template_engine import preview_path
        self.assertIsNone(preview_path("../secrets", "formal"))
        self.assertIsNone(preview_path("policy", "nope"))

    def test_logo_fetch_failure_is_silent(self):
        from backend.core.template_engine import apply_logo, load_template_shell
        doc = load_template_shell("POLICY", "modern")
        apply_logo(doc, "https://127.0.0.1:1/never-resolves.png")  # must not raise

    def test_non_http_logo_url_ignored(self):
        from unittest.mock import patch
        from backend.core.template_engine import apply_logo, load_template_shell
        doc = load_template_shell("POLICY", "modern")
        with patch("backend.core.template_engine._fetch_logo_bytes") as fetch:
            apply_logo(doc, "javascript:alert(1)")
            apply_logo(doc, "file:///etc/passwd")
            fetch.assert_not_called()


class ParagraphStyleTests(unittest.TestCase):
    def test_shell_supports_bullet_styles(self):
        """render_markdown_bullet styles bullets as List Bullet; pandoc shells
        lack it (crashed a live export)."""
        doc = load_template_shell("POLICY", "formal")
        p = doc.add_paragraph("item")
        p.style = "List Bullet"  # must not raise
        p2 = doc.add_paragraph("item2")
        p2.style = "List Number"
        self.assertIsNotNone(p.style)

    def test_shell_headers_carry_title_placeholder(self):
        """Shells must carry {{TITLE}} in the header, not baked sample titles
        (a live export shipped 'Acceptable Use Policy' on someone else's doc)."""
        doc = load_template_shell("POLICY", "formal", branding={
            "organization": "Acme", "title": "Real Title", "classification": "Internal",
        })
        header_text = " ".join(
            p.text for s in doc.sections for p in s.header.paragraphs
        )
        self.assertIn("Real Title", header_text)
        self.assertNotIn("Acceptable Use Policy", header_text)


class FrontMatterPreservationTests(unittest.TestCase):
    def test_shell_keeps_cover_and_control(self):
        """load_template_shell must preserve the designed front matter (cover
        page, Document Control, Table of Contents) — not clear the whole body."""
        doc = load_template_shell("POLICY", "formal")
        lines = [p.text.strip() for p in doc.paragraphs]
        self.assertIn("Document Control", lines)
        self.assertIn("Table of Contents", lines)
        # sample numbered sections are dropped
        self.assertFalse(any(t.startswith("1. Purpose") for t in lines))

    def test_placeholders_fill_and_none_leak(self):
        from backend.core.template_engine import fill_template_placeholders
        doc = load_template_shell("POLICY", "formal")
        fill_template_placeholders(doc, {
            "DOCUMENT_TITLE": "My Policy", "VERSION": "3.1",
            "POLICY_OWNER": "Sam Rivera", "EFFECTIVE_DATE": "2026-01-01",
        })
        joined = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells
        )
        self.assertIn("My Policy", joined)
        self.assertNotIn("{{", joined + table_text)  # nothing leaks

    def test_sample_title_removed(self):
        doc = load_template_shell("POLICY", "formal")
        joined = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("Acceptable Use Policy", joined)  # hardcoded sample gone
